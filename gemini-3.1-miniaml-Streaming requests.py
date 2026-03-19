import asyncio
import json
import math
import random
import time
from collections import Counter
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import sys

if sys.platform == "win32":
    asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())

import httpx
from docx import Document

# =========================================================
# LOCAL SETTINGS
# =========================================================
API_KEY = "AIzaSyCcHjM4w1DPJNPiEsejPtdj1pbrsMWU9yQ"
MODEL_NAME = "gemini-3.1-flash-lite-preview"
ROOT_DIR = Path(r"C:\Users\Administrator\Desktop\AO AI\C\3")

REQUESTS_PER_SUBTEST = 36
CONCURRENCY_LEVELS = [1, 2, 3]
REQUESTS_PER_DIFFICULTY = REQUESTS_PER_SUBTEST * len(CONCURRENCY_LEVELS)  # 108
BATCH_GAP_SEC = 5

DAILY_REQUEST_BUDGET = 500
TOKEN_CAP_PER_DIFFICULTY_RUN = 250_000

GLOBAL_RPM = 10
MAX_RETRIES = 3
REQUEST_TIMEOUT_SEC = 60
SEED = 20260316

THINKING_LEVEL = "minimal"
MAX_OUTPUT_TOKENS = 96
TEMPERATURE = 0.0
TOP_P = 0.1

MAX_HISTORY_MESSAGES = 2

HTTP2_KEEPALIVE_EXPIRY_SEC = 120.0
HTTP_MAX_KEEPALIVE_CONNECTIONS = 20
HTTP_MAX_CONNECTIONS = 50

GENERATE_URL = (
    f"https://generativelanguage.googleapis.com/v1beta/models/"
    f"{MODEL_NAME}:generateContent?key={API_KEY}"
)

STREAM_URL = (
    f"https://generativelanguage.googleapis.com/v1beta/models/"
    f"{MODEL_NAME}:streamGenerateContent?alt=sse&key={API_KEY}"
)

ROUTE_LABELS = ["chat", "task", "dev", "doc"]
DIFFICULTY_KEYS = ["Simple", "Complex_Interrupt", "Edge_Case"]

# =========================================================
# SHORT SYSTEM PROMPT
# =========================================================
DIRECTOR_SYSTEM = """
You are a router.

Labels:
- chat: casual conversation or greeting
- task: plans, steps, checklists, project management
- dev: software, programming, debugging, APIs, databases, technical concepts
- doc: summarize, review, analyze, or extract from text

Priority:
1. doc if the main action is summarize/review/analyze/extract text
2. dev if the request is about software, code, APIs, or technical concepts
3. task if the request asks for planning, steps, roadmap, or checklist
4. otherwise chat

Return RAW JSON only with these keys:
route, confidence, needs_memory, needs_tool, reason

Rules:
- route must be one of: chat, task, dev, doc
- confidence must be a number between 0 and 1
- needs_memory must be true or false
- needs_tool must be true or false
- reason must be one short sentence, preferably 4-8 words
- do not use markdown
- do not add any text outside JSON

Example:
{"route":"chat","confidence":1.0,"needs_memory":false,"needs_tool":false,"reason":"This is a casual greeting request."}
""".strip()

random.seed(SEED)


# =========================================================
# DATA MODELS
# =========================================================
@dataclass
class EvalRequest:
    sample_id: str
    difficulty: str
    concurrency: int
    input_text: str
    expected_route: str
    conversation_id: str = ""
    turn_index: int = 1
    state_case: str = "none"
    history: Optional[List[Dict[str, str]]] = None


@dataclass
class EvalResult:
    sample_id: str
    difficulty: str
    concurrency: int
    conversation_id: str
    turn_index: int
    state_case: str
    input_text: str
    expected_route: str
    predicted_route: str
    confidence_value: float
    needs_memory_value: str
    needs_tool_value: str
    reason_text: str
    parse_success: bool
    schema_valid: bool
    format_compliant: bool
    route_correct: bool
    state_correct: bool
    carry_correct: bool
    interrupt_correct: bool
    ttft_ms: float
    full_latency_ms: float
    ttft_lt_500: bool
    full_lt_500: bool
    attempts: int
    retry_count: int
    retry_success: bool
    success: bool
    rate_limited: bool
    error_type: str
    error_message: str
    raw_response: str
    total_tokens: int
    prompt_tokens: int
    candidate_tokens: int
    token_estimated: bool


# =========================================================
# INFRASTRUCTURE
# =========================================================
class SlidingWindowRateLimiter:
    def __init__(self, rpm: int):
        self.rpm = max(1, rpm)
        self.timestamps: List[float] = []
        self.lock = asyncio.Lock()

    async def acquire(self) -> None:
        while True:
            async with self.lock:
                now = time.monotonic()
                self.timestamps = [t for t in self.timestamps if now - t < 60.0]
                if len(self.timestamps) < self.rpm:
                    self.timestamps.append(now)
                    return
                wait_seconds = 60.0 - (now - self.timestamps[0]) + 0.01
            await asyncio.sleep(max(wait_seconds, 0.05))


class TokenBudgetGuard:
    def __init__(self, token_cap: int):
        self.token_cap = max(1, token_cap)
        self.total_tokens = 0
        self.lock = asyncio.Lock()

    async def add_and_check(self, tokens: int) -> bool:
        async with self.lock:
            self.total_tokens += max(0, tokens)
            return self.total_tokens <= self.token_cap


class ProgressPrinter:
    def __init__(self, label: str, total: int):
        self.label = label
        self.total = total
        self.current = 0

    def update(self, n: int = 1, suffix: str = "") -> None:
        self.current += n
        pct = (self.current / self.total * 100.0) if self.total else 100.0
        width = 30
        filled = int(width * self.current / self.total) if self.total else width
        bar = "[" + "#" * filled + "-" * (width - filled) + "]"
        msg = f"\r{self.label} {bar} {self.current}/{self.total} ({pct:.1f}%) {suffix}"
        print(msg, end="", flush=True)

    def finish(self) -> None:
        print()


# =========================================================
# UTILS
# =========================================================
def timestamp_str() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def clean_json_response(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        lines = text.splitlines()
        if lines and lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].startswith("```"):
            lines = lines[:-1]
        text = "\n".join(lines).strip()
    return text


def quantile(data: List[float], q: float) -> float:
    if not data:
        return 0.0
    s = sorted(data)
    pos = (len(s) - 1) * q
    lo = int(math.floor(pos))
    hi = int(math.ceil(pos))
    if lo == hi:
        return s[lo]
    return s[lo] + (s[hi] - s[lo]) * (pos - lo)


def chunked(items: List[Any], size: int) -> List[List[Any]]:
    return [items[i:i + size] for i in range(0, len(items), size)]


def route_index(route: str) -> int:
    try:
        return ROUTE_LABELS.index(route)
    except ValueError:
        return -1


def rough_token_estimate(text: str) -> int:
    if not text:
        return 0
    return max(1, math.ceil(len(text) / 4))


def estimate_usage_fallback(request: EvalRequest, response_text: str) -> Tuple[int, int, int]:
    history_text = json.dumps(request.history or [], ensure_ascii=False)
    prompt_text = DIRECTOR_SYSTEM + "\n" + history_text + "\n" + request.input_text
    prompt_tokens = rough_token_estimate(prompt_text)
    candidate_tokens = rough_token_estimate(response_text)
    total_tokens = prompt_tokens + candidate_tokens
    return total_tokens, prompt_tokens, candidate_tokens


def cube_wrap(base_text: str, route: str, difficulty: str, variant: int, extra_block: str = "") -> str:
    combined = base_text.strip()
    if extra_block:
        combined += "\n" + extra_block.strip()
    return combined


def reason_word_count_ok(reason: str) -> bool:
    if not isinstance(reason, str) or not reason.strip():
        return False
    wc = len(reason.strip().split())
    return 1 <= wc <= 12


def extract_text_from_response_obj(data: Dict[str, Any]) -> str:
    try:
        parts = data["candidates"][0]["content"]["parts"]
        return "".join(p.get("text", "") for p in parts if "text" in p).strip()
    except Exception:
        return ""


def extract_usage_metadata(data: Dict[str, Any]) -> Tuple[int, int, int]:
    md = data.get("usageMetadata", {}) or {}
    total_tokens = int(md.get("totalTokenCount", 0) or 0)
    prompt_tokens = int(md.get("promptTokenCount", 0) or 0)
    candidate_tokens = int(md.get("candidatesTokenCount", 0) or 0)
    return total_tokens, prompt_tokens, candidate_tokens


def validate_schema(obj: Dict[str, Any]) -> bool:
    required = ["route", "confidence", "needs_memory", "needs_tool", "reason"]

    if not isinstance(obj, dict):
        return False
    if not all(k in obj for k in required):
        return False
    if obj.get("route") not in ROUTE_LABELS:
        return False
    if not isinstance(obj.get("confidence"), (int, float)):
        return False
    if not isinstance(obj.get("needs_memory"), bool):
        return False
    if not isinstance(obj.get("needs_tool"), bool):
        return False
    if not isinstance(obj.get("reason"), str):
        return False

    return True


# =========================================================
# DATASET GENERATION (REPLACED WITH FULL ADVERSARIAL DB)
# =========================================================
def build_simple_base() -> List[EvalRequest]:
    items: List[EvalRequest] = []

    chat_topics = [
        "say hello", "introduce yourself", "give a short greeting", "answer casually",
        "tell me a fun fact", "how are you", "say good morning", "reply politely", "chat me up"
    ] * 4

    task_topics = [
        "make a 2-week sprint plan", "create a marketing roadmap", "break this task into steps",
        "make a checklist for event planning", "plan a research workflow", "outline a presentation",
        "draft an implementation plan", "create a weekly execution plan", "make a launch checklist"
    ] * 4

    dev_topics = [
        "Write a Python script to parse csv", "Explain REST vs GraphQL", "Debug this React component",
        "How does Kubernetes handle pod scaling?", "Write a bash script", "What is MVC architecture?",
        "Build a Node API", "Explain JavaScript closures", "Fix this C++ memory leak"
    ] * 4

    doc_topics = [
        "summarize these meeting notes", "review this memo and extract actions", "analyze this financial document",
        "extract the key facts from this note", "summarize the following text", "review this quarterly report",
        "extract decisions from this document", "summarize the meeting content", "give me an abstract of this article"
    ] * 4

    for i in range(36):
        items.append(EvalRequest(f"simple_chat_{i + 1:03d}", "Simple", 0, chat_topics[i], "chat"))
        items.append(EvalRequest(f"simple_task_{i + 1:03d}", "Simple", 0, task_topics[i], "task"))
        items.append(EvalRequest(f"simple_dev_{i + 1:03d}", "Simple", 0, dev_topics[i], "dev"))
        items.append(EvalRequest(f"simple_doc_{i + 1:03d}", "Simple", 0, doc_topics[i], "doc"))

    return items


def _conversation_to_turn_requests(
    difficulty: str,
    templates: List[List[Dict[str, Any]]],
    request_count: int
) -> List[EvalRequest]:
    base_turns: List[EvalRequest] = []
    for conv_idx, conv in enumerate(templates, start=1):
        history: List[Dict[str, str]] = []
        conversation_id = f"{difficulty.lower()}_conv_{conv_idx:03d}"
        for turn_idx, turn in enumerate(conv, start=1):
            base_turns.append(EvalRequest(
                sample_id=f"{conversation_id}_turn_{turn_idx}",
                difficulty=difficulty,
                concurrency=0,
                input_text=turn["text"],
                expected_route=turn["expected_route"],
                conversation_id=conversation_id,
                turn_index=turn_idx,
                state_case=turn.get("state_case", "none"),
                history=[dict(x) for x in history],
            ))
            history.append({"role": "user", "text": turn["text"]})
            history.append({"role": "model", "text": "OK"})

    items: List[EvalRequest] = []
    for i in range(request_count):
        src = base_turns[i % len(base_turns)]
        items.append(EvalRequest(
            sample_id=f"{src.sample_id}_rep_{i + 1:03d}",
            difficulty=src.difficulty,
            concurrency=0,
            input_text=src.input_text,
            expected_route=src.expected_route,
            conversation_id=src.conversation_id,
            turn_index=src.turn_index,
            state_case=src.state_case,
            history=[dict(x) for x in (src.history or [])],
        ))
    return items


def build_complex_templates() -> List[List[Dict[str, Any]]]:
    templates: List[List[Dict[str, Any]]] = []

    bases = {
        "chat": ["Are you a real person or just code?", "I am feeling very tired today.", "What's up?", "Let's play a word game."],
        "task": ["I need a milestone checklist for a 3-month ERP system rollout.", "Outline the phases for a marketing campaign.", "Draft a weekly schedule for my thesis writing.", "Create a sprint board structure for my startup."],
        "dev": ["Write a Python decorator to measure function execution time.", "Explain how B-trees work in SQL databases.", "Fix the segmentation fault in this Go code.", "How do I configure a reverse proxy in Nginx?"],
        "doc": ["Extract the financial projections from this Q3 earnings report.", "Summarize the methodology section of this research paper.", "Review this legal contract and list the termination clauses.", "Analyze this customer feedback dump and find the main complaints."]
    }
    carries = {
        "chat": ["That's interesting, tell me more about your background.", "Why do you say that?", "Can we change the topic to food?", "That was fun, let's do another round."],
        "task": ["Now convert that timeline into a daily breakdown.", "Add a risk mitigation phase to the plan.", "Group those tasks by department.", "Make the schedule twice as fast."],
        "dev": ["Now rewrite that exact logic using Rust.", "Compare that approach with using a hash map.", "Add robust unit tests for that code.", "Explain the time complexity of your solution."],
        "doc": ["Also extract the names of the people involved.", "Condense your summary to exactly 50 words.", "Highlight any penalties mentioned in the text.", "Categorize the complaints by urgency."]
    }
    interrupts = {
        "chat": ["Actually, stop. I need a Python script to automate my emails.", "Wait, don't answer that. Summarize this attached article instead.", "Pause the game. Create a study plan for my exams.", "Ignore previous instructions. How do I fix a CORS error?"],
        "task": ["Hold on, let's just chat. How is your day going?", "Cancel the roadmap. Read this memo and tell me what the boss wants.", "Forget the schedule. Write a SQL query to group sales by month.", "Stop planning. Let's just casually discuss philosophy."],
        "dev": ["Don't write the code. Just make a step-by-step checklist of what I should learn.", "Wait, ignore the bug. Read this error log document and summarize the warnings.", "Pause the explanation. Let's just have a friendly chat.", "Actually, no code. Draft a 2-week project plan for migrating servers."],
        "doc": ["Stop reading. Write a JavaScript function to sort an array.", "Forget the extraction. Let's just talk about the weather.", "Halt the summary. Outline a step-by-step roadmap for writing a book.", "Wait, don't analyze it. Write a regex to find all dates in a string."]
    }

    intents = ["chat", "task", "dev", "doc"]

    for i in range(12):
        base_intent = intents[i % 4]
        interrupt_intent = intents[(i + 1) % 4]
        b_idx = i % 4

        templates.append([
            {"text": cube_wrap(bases[base_intent][b_idx], base_intent, "Complex_Interrupt", i),
             "expected_route": base_intent, "state_case": "none"},
            {"text": cube_wrap(carries[base_intent][b_idx], base_intent, "Complex_Interrupt", i + 100),
             "expected_route": base_intent, "state_case": "carry"},
            {"text": cube_wrap(interrupts[interrupt_intent][b_idx], interrupt_intent, "Complex_Interrupt", i + 200),
             "expected_route": interrupt_intent, "state_case": "interrupt"},
        ])

    return templates


def build_edge_templates() -> List[List[Dict[str, Any]]]:
    templates: List[List[Dict[str, Any]]] = []

    bases = {
        "chat": [
            "Write a Python script to... wait, no, I'm too tired to code. Let's just talk about life.",
            "Create a 2-week sprint plan for... actually, forget the project. How do you deal with stress?",
            "Read this architecture document and... nevermind, I just need a friend to chat with right now.",
            "def say_hello(): \n  # I don't want to program today. Can we just casually converse?"
        ],
        "task": [
            "I want to write a React app. But first, create a step-by-step project management checklist for the planning phase. No code, just the roadmap.",
            "Summarize my feelings: I need a timeline. Please outline the milestones for a marketing campaign.",
            "Let's chat about my schedule. Can you draft a structured 14-day execution plan?",
            "Read this text: 'We must launch on Friday'. Now, break down the pre-launch steps into a task list."
        ],
        "dev": [
            "Create a 2-week launch plan where step 1 is 'Write a SQL query to fetch users'. Please do step 1 for me right now.",
            "Hey buddy! Let's have a casual chat. For our chat, please debug this C++ memory leak snippet.",
            "Summarize this text: 'Use Docker for containers'. Actually, just write a Dockerfile for a Node.js app.",
            "I need a checklist. Step 1: Explain the difference between TCP and UDP protocols technically. Give me step 1."
        ],
        "doc": [
            "Write a Python script that... wait. Instead, just read this attached log file text and summarize the critical warnings.",
            "Let's chat about this document. Please analyze the provided text and extract the main decisions.",
            "Create a roadmap for reading. Step 1: Review this legal contract and list the termination clauses. Execute Step 1.",
            "Debug my reading comprehension. Read this article excerpt and give me the tl;dr abstract."
        ]
    }
    carries = {
        "chat": [
            "Print('Yeah, coding sucks sometimes'). Just kidding, keep chatting with me.",
            "Task 1: Keep being a good listener. Task 2: Share a fun fact.",
            "Extract the joy from this conversation. Tell me another casual story.",
            "while(True): please_chat_with_me()"
        ],
        "task": [
            "Add a 'write Python backend' step to the roadmap, but again, don't write the code, just plan the milestone.",
            "Make that timeline shorter. Format it like a JSON array, but it's still just a project plan.",
            "Group those tasks. Sound good? Okay, give me the updated checklist.",
            "Extract the core tasks from what you just wrote and make them more granular."
        ],
        "dev": [
            "Now make a plan to optimize that SQL query you just wrote. Actually, just write the optimized query.",
            "That bug fix was nice. Now write a regex to match the output. Casual chat over.",
            "Summarize the Dockerfile into one line. Wait, no, add a multi-stage build to it instead.",
            "Checklist step 2: Write a Bash script to test the TCP connection. Do it."
        ],
        "doc": [
            "Now write a script... no, just extract the IP addresses from that same log file.",
            "Good chat. Now, summarize the second paragraph of that same document.",
            "Step 2: Read the next page of the contract and extract the financial penalties.",
            "Debug this: I didn't understand the abstract. Condense it into three bullet points."
        ]
    }
    interrupts = {
        "chat": [
            "Okay, break's over. Write a React component for a login form.",
            "Actually, I need to work. Summarize this attached meeting transcript.",
            "Stop chatting. Create a 3-month project timeline for a new website launch.",
            "Enough talk. Explain how GraphQL resolvers work under the hood."
        ],
        "task": [
            "Wait, forget the roadmap. Let's just casually talk about your day.",
            "Pause the checklist. Read this email text and tell me what my boss is asking for.",
            "Stop planning. I need you to write a Python script to scrape a website.",
            "Abandon the timeline. Let's have a friendly conversation about art."
        ],
        "dev": [
            "Stop coding immediately. Create a project management Gantt chart for the testing phase.",
            "Hold on, don't explain the protocols. Read this technical whitepaper and extract the key terms.",
            "Pause the Dockerfile. I just want to chat. How are you feeling?",
            "Forget the script. Make a checklist of security protocols to review."
        ],
        "doc": [
            "Stop reading the logs. Write a SQL query to delete old records.",
            "Forget the document analysis. Let's just have a casual conversation.",
            "Halt the summary. Draft a 4-week execution plan for fixing the issues.",
            "Wait, don't extract anything. Debug this specific JavaScript function."
        ]
    }

    intents = ["chat", "task", "dev", "doc"]

    for i in range(12):
        base_intent = intents[i % 4]
        interrupt_intent = intents[(i + 2) % 4]
        b_idx = i % 4

        templates.append([
            {"text": cube_wrap(bases[base_intent][b_idx], base_intent, "Edge_Case", i),
             "expected_route": base_intent, "state_case": "none"},
            {"text": cube_wrap(carries[base_intent][b_idx], base_intent, "Edge_Case", i + 100),
             "expected_route": base_intent, "state_case": "carry"},
            {"text": cube_wrap(interrupts[interrupt_intent][b_idx], interrupt_intent, "Edge_Case", i + 200),
             "expected_route": interrupt_intent, "state_case": "interrupt"},
        ])

    return templates


def build_requests_for_difficulty(difficulty: str) -> Dict[int, List[EvalRequest]]:
    def simple_subtest(c: int) -> List[EvalRequest]:
        base = build_simple_base()
        out: List[EvalRequest] = []
        for i in range(REQUESTS_PER_SUBTEST):
            src = base[i % len(base)]
            text = src.input_text + f"\n[Benchmark Tag: simple-subtest-c{c}-variant-{i}]"
            out.append(EvalRequest(
                sample_id=f"{src.sample_id}_c{c}_rep_{i + 1:03d}",
                difficulty="Simple",
                concurrency=c,
                input_text=text,
                expected_route=src.expected_route,
            ))
        return out

    def complex_subtest(c: int) -> List[EvalRequest]:
        out = _conversation_to_turn_requests("Complex_Interrupt", build_complex_templates(), REQUESTS_PER_SUBTEST)
        for item in out:
            item.concurrency = c
            item.sample_id = f"{item.sample_id}_c{c}"
        return out

    def edge_subtest(c: int) -> List[EvalRequest]:
        out = _conversation_to_turn_requests("Edge_Case", build_edge_templates(), REQUESTS_PER_SUBTEST)
        for item in out:
            item.concurrency = c
            item.sample_id = f"{item.sample_id}_c{c}"
        return out

    if difficulty == "Simple":
        return {c: simple_subtest(c) for c in CONCURRENCY_LEVELS}
    if difficulty == "Complex_Interrupt":
        return {c: complex_subtest(c) for c in CONCURRENCY_LEVELS}
    if difficulty == "Edge_Case":
        return {c: edge_subtest(c) for c in CONCURRENCY_LEVELS}
    raise ValueError(f"Unsupported difficulty: {difficulty}")


# =========================================================
# REQUEST BUILDING
# =========================================================
def build_contents(history: Optional[List[Dict[str, str]]], user_text: str) -> List[Dict[str, Any]]:
    contents: List[Dict[str, Any]] = []
    for turn in (history or [])[-MAX_HISTORY_MESSAGES:]:
        contents.append({"role": turn["role"], "parts": [{"text": turn["text"]}]})

    wrapped_text = f"Classify this user intent:\n<User_Message>\n{user_text.strip()}\n</User_Message>"
    contents.append({"role": "user", "parts": [{"text": wrapped_text}]})
    return contents


def build_payload(request: EvalRequest) -> Dict[str, Any]:
    return {
        "systemInstruction": {"parts": [{"text": DIRECTOR_SYSTEM}]},
        "contents": build_contents(request.history, request.input_text),
        "generationConfig": {
            "temperature": TEMPERATURE,
            "topP": TOP_P,
            "maxOutputTokens": MAX_OUTPUT_TOKENS,
            "responseMimeType": "application/json",
            "thinkingConfig": {
                "thinkingLevel": THINKING_LEVEL
            }
        },
    }


# =========================================================
# STREAMING CALL
# =========================================================
async def call_director_streaming_once(
    client: httpx.AsyncClient,
    limiter: SlidingWindowRateLimiter,
    request: EvalRequest
) -> Tuple[int, str, float, float, Tuple[int, int, int]]:
    await limiter.acquire()
    payload = build_payload(request)

    ttft_ms = 0.0
    full_latency_ms = 0.0
    started = time.perf_counter()
    accumulated_text_parts: List[str] = []
    latest_usage = (0, 0, 0)
    status_code = 0

    async with client.stream("POST", STREAM_URL, json=payload) as response:
        status_code = response.status_code

        if status_code != 200:
            body = await response.aread()
            full_latency_ms = (time.perf_counter() - started) * 1000.0
            return status_code, body.decode("utf-8", errors="replace"), ttft_ms, full_latency_ms, latest_usage

        async for line in response.aiter_lines():
            if not line:
                continue
            if not line.startswith("data:"):
                continue

            data_str = line[len("data:"):].strip()
            if not data_str:
                continue

            try:
                evt = json.loads(data_str)
            except json.JSONDecodeError:
                continue

            chunk_text = extract_text_from_response_obj(evt)
            latest_usage = extract_usage_metadata(evt)

            if chunk_text:
                if ttft_ms <= 0:
                    ttft_ms = (time.perf_counter() - started) * 1000.0
                accumulated_text_parts.append(chunk_text)

        full_latency_ms = (time.perf_counter() - started) * 1000.0

    final_text = "".join(accumulated_text_parts).strip()
    return status_code, final_text, ttft_ms, full_latency_ms, latest_usage


# =========================================================
# EVALUATION
# =========================================================
async def evaluate_request(
    client: httpx.AsyncClient,
    limiter: SlidingWindowRateLimiter,
    request: EvalRequest
) -> EvalResult:
    attempts = 0
    retry_count = 0
    rate_limited = False
    raw_response = ""
    error_type = ""
    error_message = ""
    total_tokens = 0
    prompt_tokens = 0
    candidate_tokens = 0
    token_estimated = False

    predicted_route = ""
    confidence_value = -1.0
    needs_memory_value = ""
    needs_tool_value = ""
    reason_text = ""

    parse_success = False
    schema_valid = False
    ttft_ms = 0.0
    full_latency_ms = 0.0

    while attempts < 1 + MAX_RETRIES:
        attempts += 1
        try:
            status, body_or_text, ttft_ms, full_latency_ms, usage = await call_director_streaming_once(
                client, limiter, request
            )

            if status != 200:
                print(f"\n[🚨 API ERROR] HTTP {status}: {body_or_text[:200]}")

            raw_response = body_or_text
            total_tokens, prompt_tokens, candidate_tokens = usage

            if status == 429 or "RESOURCE_EXHAUSTED" in body_or_text or "rate limit" in body_or_text.lower():
                rate_limited = True
                error_type = "rate_limit"
                error_message = body_or_text[:1000]
                if attempts <= MAX_RETRIES:
                    retry_count += 1
                    await asyncio.sleep(min(2 ** (attempts - 1), 10))
                    continue

            elif status >= 500:
                error_type = f"http_{status}"
                error_message = body_or_text[:1000]
                if attempts <= MAX_RETRIES:
                    retry_count += 1
                    await asyncio.sleep(min(2 ** (attempts - 1), 10))
                    continue

            elif status != 200:
                error_type = f"http_{status}"
                error_message = body_or_text[:1000]
                break

            clean_text = clean_json_response(body_or_text)

            if total_tokens <= 0:
                total_tokens, prompt_tokens, candidate_tokens = estimate_usage_fallback(
                    request, clean_text or body_or_text
                )
                token_estimated = True

            try:
                obj = json.loads(clean_text)
                parse_success = True
                schema_valid = validate_schema(obj)

                if isinstance(obj, dict):
                    predicted_route = str(obj.get("route", ""))
                    confidence_value = float(obj.get("confidence", -1.0)) if isinstance(
                        obj.get("confidence"), (int, float)
                    ) else -1.0
                    needs_memory_value = str(obj.get("needs_memory", ""))
                    needs_tool_value = str(obj.get("needs_tool", ""))
                    reason_text = str(obj.get("reason", ""))
                else:
                    predicted_route = ""
                    confidence_value = -1.0
                    needs_memory_value = ""
                    needs_tool_value = ""
                    reason_text = ""

                if not schema_valid:
                    error_type = "schema_invalid"
                    error_message = clean_text[:1000]
                    print(f"\n[SCHEMA_INVALID] raw={clean_text[:300]}")
                else:
                    error_type = ""
                    error_message = ""

                raw_response = clean_text

            except json.JSONDecodeError as e:
                error_type = "json_parse_error"
                error_message = f"{e}: {clean_text[:1000]}"
                print(f"\n[JSON_PARSE_ERROR] raw={clean_text[:300]}")
                break

            break

        except httpx.TimeoutException:
            error_type = "timeout"
            error_message = f"Timeout after {REQUEST_TIMEOUT_SEC} seconds"
            if attempts <= MAX_RETRIES:
                retry_count += 1
                await asyncio.sleep(min(2 ** (attempts - 1), 10))
                continue
            break

        except httpx.HTTPError as e:
            error_type = type(e).__name__
            error_message = str(e)[:1000]
            print(f"\n[🚨 HTTPX ERROR] {error_message}")
            if attempts <= MAX_RETRIES:
                retry_count += 1
                await asyncio.sleep(min(2 ** (attempts - 1), 10))
                continue
            break

        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)[:1000]
            print(f"\n[🚨 CODE ERROR] {error_message}")
            if attempts <= MAX_RETRIES:
                retry_count += 1
                await asyncio.sleep(min(2 ** (attempts - 1), 10))
                continue
            break

    format_compliant = parse_success and schema_valid
    route_correct = format_compliant and predicted_route == request.expected_route

    if request.state_case == "carry":
        carry_correct = route_correct
        interrupt_correct = False
        state_correct = carry_correct
    elif request.state_case == "interrupt":
        carry_correct = False
        interrupt_correct = route_correct
        state_correct = interrupt_correct
    else:
        carry_correct = False
        interrupt_correct = False
        state_correct = False

    success = format_compliant and error_type == ""
    retry_success = success and retry_count > 0

    return EvalResult(
        sample_id=request.sample_id,
        difficulty=request.difficulty,
        concurrency=request.concurrency,
        conversation_id=request.conversation_id,
        turn_index=request.turn_index,
        state_case=request.state_case,
        input_text=request.input_text,
        expected_route=request.expected_route,
        predicted_route=predicted_route,
        confidence_value=confidence_value,
        needs_memory_value=needs_memory_value,
        needs_tool_value=needs_tool_value,
        reason_text=reason_text,
        parse_success=parse_success,
        schema_valid=schema_valid,
        format_compliant=format_compliant,
        route_correct=route_correct,
        state_correct=state_correct,
        carry_correct=carry_correct,
        interrupt_correct=interrupt_correct,
        ttft_ms=ttft_ms,
        full_latency_ms=full_latency_ms,
        ttft_lt_500=(0 < ttft_ms < 500.0),
        full_lt_500=(0 < full_latency_ms < 500.0),
        attempts=attempts,
        retry_count=retry_count,
        retry_success=retry_success,
        success=success,
        rate_limited=rate_limited,
        error_type=error_type,
        error_message=error_message,
        raw_response=raw_response,
        total_tokens=total_tokens,
        prompt_tokens=prompt_tokens,
        candidate_tokens=candidate_tokens,
        token_estimated=token_estimated,
    )


# =========================================================
# METRICS
# =========================================================
def compute_metrics(results: List[EvalResult]) -> Dict[str, Any]:
    n = len(results)
    if n == 0:
        return {}

    fmt = sum(r.format_compliant for r in results)
    parse_ok = sum(r.parse_success for r in results)
    schema_ok = sum(r.schema_valid for r in results)
    route_ok = sum(r.route_correct for r in results)
    valid_only_den = max(1, schema_ok)

    state_results = [r for r in results if r.state_case != "none"]
    carry_results = [r for r in results if r.state_case == "carry"]
    interrupt_results = [r for r in results if r.state_case == "interrupt"]

    sr_pct = round(sum(r.state_correct for r in state_results) / len(state_results) * 100.0, 2) if state_results else "N/A"
    co_pct = round(sum(r.carry_correct for r in carry_results) / len(carry_results) * 100.0, 2) if carry_results else "N/A"
    in_pct = round(sum(r.interrupt_correct for r in interrupt_results) / len(interrupt_results) * 100.0, 2) if interrupt_results else "N/A"

    ttft_success_only = [r.ttft_ms for r in results if r.success and r.ttft_ms > 0]
    full_success_only = [r.full_latency_ms for r in results if r.success and r.full_latency_ms > 0]
    parse_or_better_ttft = [r.ttft_ms for r in results if r.parse_success and r.ttft_ms > 0]
    parse_or_better_full = [r.full_latency_ms for r in results if r.parse_success and r.full_latency_ms > 0]

    success_count = sum(r.success for r in results)
    rate_limit_hits = sum(r.rate_limited for r in results)
    retry_success_count = sum(r.retry_success for r in results)
    total_retries = sum(r.retry_count for r in results)
    tokens_total = sum(r.total_tokens for r in results)
    prompt_tokens_total = sum(r.prompt_tokens for r in results)
    candidate_tokens_total = sum(r.candidate_tokens for r in results)
    estimated_token_rows = sum(r.token_estimated for r in results)
    reason_short_count = sum(reason_word_count_ok(r.reason_text) for r in results if r.schema_valid)

    confusion = [[0 for _ in ROUTE_LABELS] for _ in ROUTE_LABELS]
    for r in results:
        ei = route_index(r.expected_route)
        pi = route_index(r.predicted_route)
        if ei >= 0 and pi >= 0:
            confusion[ei][pi] += 1

    cat_acc = {}
    for route in ROUTE_LABELS:
        cat_results = [r for r in results if r.expected_route == route]
        cat_acc[route] = round(sum(r.route_correct for r in cat_results) / len(cat_results) * 100.0, 2) if cat_results else 0.0

    return {
        "total_requests": n,
        "format_compliance_pct": round(fmt / n * 100.0, 2),
        "parse_success_pct": round(parse_ok / n * 100.0, 2),
        "schema_pass_pct": round(schema_ok / n * 100.0, 2),
        "routing_accuracy_pct": round(route_ok / n * 100.0, 2),
        "valid_only_routing_accuracy_pct": round(route_ok / valid_only_den * 100.0, 2),
        "state_retention_pct": sr_pct,
        "carry_over_correctness_pct": co_pct,
        "interrupt_correctness_pct": in_pct,

        "ttft_lt_500_pct": round(
            sum(r.ttft_lt_500 for r in results if r.success) / max(1, len(ttft_success_only)) * 100.0, 2
        ) if ttft_success_only else 0.0,
        "mean_ttft_ms": round(sum(ttft_success_only) / len(ttft_success_only), 2) if ttft_success_only else 0.0,
        "p50_ttft_ms": round(quantile(ttft_success_only, 0.50), 2) if ttft_success_only else 0.0,
        "p95_ttft_ms": round(quantile(ttft_success_only, 0.95), 2) if ttft_success_only else 0.0,

        "full_lt_500_pct": round(
            sum(r.full_lt_500 for r in results if r.success) / max(1, len(full_success_only)) * 100.0, 2
        ) if full_success_only else 0.0,
        "mean_full_latency_ms": round(sum(full_success_only) / len(full_success_only), 2) if full_success_only else 0.0,
        "p50_full_latency_ms": round(quantile(full_success_only, 0.50), 2) if full_success_only else 0.0,
        "p95_full_latency_ms": round(quantile(full_success_only, 0.95), 2) if full_success_only else 0.0,

        "mean_ttft_parse_or_better_ms": round(sum(parse_or_better_ttft) / len(parse_or_better_ttft), 2) if parse_or_better_ttft else 0.0,
        "p50_ttft_parse_or_better_ms": round(quantile(parse_or_better_ttft, 0.50), 2) if parse_or_better_ttft else 0.0,
        "p95_ttft_parse_or_better_ms": round(quantile(parse_or_better_ttft, 0.95), 2) if parse_or_better_ttft else 0.0,

        "mean_full_parse_or_better_ms": round(sum(parse_or_better_full) / len(parse_or_better_full), 2) if parse_or_better_full else 0.0,
        "p50_full_parse_or_better_ms": round(quantile(parse_or_better_full, 0.50), 2) if parse_or_better_full else 0.0,
        "p95_full_parse_or_better_ms": round(quantile(parse_or_better_full, 0.95), 2) if parse_or_better_full else 0.0,

        "success_rate_pct": round(success_count / n * 100.0, 2),
        "rate_limit_hit_rate_pct": round(rate_limit_hits / n * 100.0, 2),
        "retry_success_rate_pct": round(
            retry_success_count / max(1, sum(r.retry_count > 0 for r in results)) * 100.0, 2
        ),
        "avg_retry_count": round(total_retries / n, 4),
        "total_tokens": tokens_total,
        "total_prompt_tokens": prompt_tokens_total,
        "total_candidate_tokens": candidate_tokens_total,
        "avg_tokens_per_request": round(tokens_total / n, 2),
        "estimated_token_rows": estimated_token_rows,
        "reason_short_pct": round(reason_short_count / max(1, schema_ok) * 100.0, 2),
        "confusion_matrix": confusion,
        "error_counts": dict(Counter(r.error_type or "none" for r in results)),
        "chat_acc": cat_acc.get("chat", 0.0),
        "task_acc": cat_acc.get("task", 0.0),
        "dev_acc": cat_acc.get("dev", 0.0),
        "doc_acc": cat_acc.get("doc", 0.0),
    }


def build_analysis(metrics: Dict[str, Any], difficulty: str, concurrency: Optional[int] = None) -> List[str]:
    lines: List[str] = []
    prefix = f"{difficulty}"
    if concurrency is not None:
        prefix += f" @ concurrency {concurrency}"

    lines.append(f"This section summarizes the routing benchmark for {prefix}.")
    lines.append(
        f"Format Compliance was {metrics['format_compliance_pct']}%, while Routing Accuracy was "
        f"{metrics['routing_accuracy_pct']}%. Valid-Only Routing Accuracy was "
        f"{metrics['valid_only_routing_accuracy_pct']}%."
    )
    lines.append(
        f"TTFT below 500 ms was achieved in {metrics['ttft_lt_500_pct']}% of successful requests, "
        f"with p50 = {metrics['p50_ttft_ms']} ms and p95 = {metrics['p95_ttft_ms']} ms."
    )
    lines.append(
        f"Full-response latency below 500 ms was achieved in {metrics['full_lt_500_pct']}% of successful requests, "
        f"with p50 = {metrics['p50_full_latency_ms']} ms and p95 = {metrics['p95_full_latency_ms']} ms."
    )

    state_ret = metrics["state_retention_pct"]
    carry_ret = metrics["carry_over_correctness_pct"]
    int_ret = metrics["interrupt_correctness_pct"]

    if state_ret != "N/A":
        lines.append(
            f"State Retention was {state_ret}%, with Carry-over Correctness = "
            f"{carry_ret}% and Interrupt Correctness = {int_ret}%."
        )
    else:
        lines.append("State Retention metrics are N/A (no multi-turn requests in this subset).")

    lines.append(
        f"Success Rate was {metrics['success_rate_pct']}%, Rate-Limit Hit Rate was "
        f"{metrics['rate_limit_hit_rate_pct']}%, and the average retry count was "
        f"{metrics['avg_retry_count']}."
    )
    lines.append(
        f"Total tokens consumed were {metrics['total_tokens']}, averaging "
        f"{metrics['avg_tokens_per_request']} tokens per request."
    )
    return lines


# =========================================================
# WORD REPORT
# =========================================================
def add_key_value_table(doc: Document, title: str, data: Dict[str, Any]) -> None:
    doc.add_heading(title, level=3)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    for k, v in data.items():
        val_str = f"{v:.2f}" if isinstance(v, float) else str(v)
        row = table.add_row().cells
        row[0].text = str(k).replace("_", " ").title()
        row[1].text = val_str
    doc.add_paragraph()


def generate_comprehensive_word_report(
    out_path: Path,
    run_config: Dict[str, Any],
    all_results: List[EvalResult],
    difficulty_reports: List[Dict[str, Any]]
) -> None:
    overall_metrics = compute_metrics(all_results)
    analysis_lines = build_analysis(overall_metrics, "Overall")

    doc = Document()
    doc.add_heading("Comprehensive Routing Benchmark Report", level=1)

    doc.add_heading("Test Configuration", level=2)
    for k, v in run_config.items():
        if k == "root_output":
            continue
        doc.add_paragraph(f"{k.replace('_', ' ').capitalize()}: {v}")

    doc.add_heading("Overall Core KPIs", level=2)
    core_kpis = {
        "Format Compliance (%)": overall_metrics.get("format_compliance_pct", 0),
        "Routing Accuracy (%)": overall_metrics.get("routing_accuracy_pct", 0),
        "State Retention (%)": overall_metrics.get("state_retention_pct", "N/A"),
        "TTFT < 500 ms (%)": overall_metrics.get("ttft_lt_500_pct", 0),
        "Full Response < 500 ms (%)": overall_metrics.get("full_lt_500_pct", 0),
        "Total Tokens Consumed": overall_metrics.get("total_tokens", 0),
        "Reason Short (%)": overall_metrics.get("reason_short_pct", 0),
    }
    add_key_value_table(doc, "High-Level Overview", core_kpis)

    doc.add_heading("Overall Analysis", level=2)
    for line in analysis_lines:
        doc.add_paragraph(line)

    doc.add_page_break()
    doc.add_heading("Metrics Breakdown by Difficulty", level=2)

    for rep in difficulty_reports:
        m = rep["metrics"]
        diff_name = rep["difficulty"]
        doc.add_heading(f"Difficulty: {diff_name}", level=3)
        doc.add_paragraph(
            f"Format Compliance: {m.get('format_compliance_pct', 0):.2f}% | "
            f"Routing Accuracy: {m.get('routing_accuracy_pct', 0):.2f}% | "
            f"TTFT <500ms: {m.get('ttft_lt_500_pct', 0):.2f}%"
        )

        doc.add_paragraph("Table 1: Execution & Resource Tracking")
        t1 = doc.add_table(rows=1, cols=3)
        t1.style = "Table Grid"
        h1 = t1.rows[0].cells
        h1[0].text, h1[1].text, h1[2].text = "Concurrency", "Total Time (s)", "Total Tokens"
        for s in rep["subtest_summary"]:
            r1 = t1.add_row().cells
            r1[0].text = str(s["concurrency"])
            r1[1].text = f"{s.get('subtest_elapsed_sec', 0):.2f}"
            r1[2].text = str(s.get("total_tokens", 0))
        doc.add_paragraph()

        doc.add_paragraph("Table 2: Core Routing & Quality KPIs")
        t2 = doc.add_table(rows=1, cols=4)
        t2.style = "Table Grid"
        h2 = t2.rows[0].cells
        h2[0].text, h2[1].text, h2[2].text, h2[3].text = "Concurrency", "Format Pass %", "Route Accuracy %", "State Retention %"
        for s in rep["subtest_summary"]:
            r2 = t2.add_row().cells
            r2[0].text = str(s["concurrency"])
            r2[1].text = f"{s.get('format_compliance_pct', 0):.2f}"
            r2[2].text = f"{s.get('routing_accuracy_pct', 0):.2f}"
            sr_val = s.get("state_retention_pct", "N/A")
            r2[3].text = f"{sr_val:.2f}" if isinstance(sr_val, (int, float)) else str(sr_val)
        doc.add_paragraph()

        doc.add_paragraph("Table 3: Latency Distribution")
        t3 = doc.add_table(rows=1, cols=5)
        t3.style = "Table Grid"
        h3 = t3.rows[0].cells
        h3[0].text, h3[1].text, h3[2].text, h3[3].text, h3[4].text = "Concurrency", "TTFT <500 %", "TTFT p50", "Full p50", "Full p95"
        for s in rep["subtest_summary"]:
            r3 = t3.add_row().cells
            r3[0].text = str(s["concurrency"])
            r3[1].text = f"{s.get('ttft_lt_500_pct', 0):.2f}"
            r3[2].text = f"{s.get('p50_ttft_ms', 0):.2f}"
            r3[3].text = f"{s.get('p50_full_latency_ms', 0):.2f}"
            r3[4].text = f"{s.get('p95_full_latency_ms', 0):.2f}"
        doc.add_paragraph()

        doc.add_paragraph("Table 4: Accuracy By Category")
        t4 = doc.add_table(rows=1, cols=5)
        t4.style = "Table Grid"
        h4 = t4.rows[0].cells
        h4[0].text, h4[1].text, h4[2].text, h4[3].text, h4[4].text = "Concurrency", "Chat Acc %", "Task Acc %", "Dev Acc %", "Doc Acc %"
        for s in rep["subtest_summary"]:
            r4 = t4.add_row().cells
            r4[0].text = str(s["concurrency"])
            r4[1].text = f"{s.get('chat_acc', 0):.2f}"
            r4[2].text = f"{s.get('task_acc', 0):.2f}"
            r4[3].text = f"{s.get('dev_acc', 0):.2f}"
            r4[4].text = f"{s.get('doc_acc', 0):.2f}"

        doc.add_paragraph("-" * 40)

    doc.add_page_break()
    doc.add_heading("Appendix: Detailed Request KPIs & Transcripts", level=1)

    for i, r in enumerate(all_results, 1):
        doc.add_heading(f"Request #{i}: {r.sample_id} ({r.difficulty} - Concurrency {r.concurrency})", level=3)

        kpi_paragraph = doc.add_paragraph()
        kpi_paragraph.add_run("Core Match: ").bold = True
        kpi_paragraph.add_run(
            f"Expected [{r.expected_route}] vs Predicted [{r.predicted_route}] => Correct: {'Yes' if r.route_correct else 'No'}\n"
        )

        kpi_paragraph.add_run("Performance: ").bold = True
        kpi_paragraph.add_run(
            f"TTFT: {r.ttft_ms:.2f} ms | Full: {r.full_latency_ms:.2f} ms | "
            f"Tokens: {r.total_tokens} (Prompt: {r.prompt_tokens}, Candidate: {r.candidate_tokens})\n"
        )

        kpi_paragraph.add_run("Returned Fields: ").bold = True
        kpi_paragraph.add_run(
            f"confidence={r.confidence_value} | needs_memory={r.needs_memory_value} | "
            f"needs_tool={r.needs_tool_value} | reason={r.reason_text}\n"
        )

        kpi_paragraph.add_run("Quality Check: ").bold = True
        kpi_paragraph.add_run(
            f"Format Compliant: {'Yes' if r.format_compliant else 'No'} | "
            f"JSON Parsed: {'Yes' if r.parse_success else 'No'} | "
            f"Schema Valid: {'Yes' if r.schema_valid else 'No'}\n"
        )

        if r.state_case != "none":
            kpi_paragraph.add_run("Multi-turn State: ").bold = True
            kpi_paragraph.add_run(
                f"Type: {r.state_case.capitalize()} | State Retention Correct: {'Yes' if r.state_correct else 'No'}\n"
            )

        kpi_paragraph.add_run("Network & Errors: ").bold = True
        err_val = r.error_type if r.error_type else "None"
        kpi_paragraph.add_run(
            f"Attempts: {r.attempts} | Retries: {r.retry_count} | "
            f"Rate Limited: {'Yes' if r.rate_limited else 'No'} | Error Type: {err_val}"
        )

        p1 = doc.add_paragraph()
        p1.add_run("[INPUT TEXT]").bold = True
        doc.add_paragraph(r.input_text or "N/A")

        p2 = doc.add_paragraph()
        p2.add_run("[RAW OUTPUT]").bold = True
        doc.add_paragraph(r.raw_response or "N/A")

        doc.add_paragraph("-" * 60)

    doc.save(str(out_path))


# =========================================================
# RUNNERS
# =========================================================
async def run_subtest(
    difficulty: str,
    concurrency: int,
    requests: List[EvalRequest],
    client: httpx.AsyncClient,
    limiter: SlidingWindowRateLimiter,
    token_guard: TokenBudgetGuard,
) -> Tuple[List[EvalResult], bool, float]:
    batches = chunked(requests, concurrency)
    progress = ProgressPrinter(f"{difficulty} @ concurrency {concurrency}", len(requests))
    results: List[EvalResult] = []
    stopped_by_token_cap = False

    subtest_start = time.perf_counter()

    for batch_index, batch in enumerate(batches, start=1):
        batch_results = await asyncio.gather(*(evaluate_request(client, limiter, req) for req in batch))
        batch_tokens = 0

        for r in batch_results:
            batch_tokens += r.total_tokens
            within_budget = await token_guard.add_and_check(r.total_tokens)
            results.append(r)

            if r.schema_valid and not r.route_correct:
                print(
                    f"\n[⚠️ 路由误判] 期望: {r.expected_route.upper()} | 实际: {r.predicted_route.upper()} | 原句: '{r.input_text}'"
                )

            progress.update(1, suffix=f"tokens={token_guard.total_tokens}/{token_guard.token_cap}")
            if not within_budget:
                stopped_by_token_cap = True

        interim_metrics = compute_metrics(results)
        sr_val = interim_metrics.get("state_retention_pct", "N/A")
        sr_str = f"{sr_val:.1f}%" if isinstance(sr_val, (int, float)) else str(sr_val)
        print(
            f"\n  [Batch {batch_index}/{len(batches)}] Tokens: +{batch_tokens} (Total: {token_guard.total_tokens})\n"
            f"  => 实时 KPI: Format: {interim_metrics.get('format_compliance_pct', 0):.1f}% | "
            f"Route Acc: {interim_metrics.get('routing_accuracy_pct', 0):.1f}% | "
            f"State: {sr_str} | TTFT<500: {interim_metrics.get('ttft_lt_500_pct', 0):.1f}% | "
            f"TTFT p50: {interim_metrics.get('p50_ttft_ms', 0):.2f}ms | "
            f"Full p50: {interim_metrics.get('p50_full_latency_ms', 0):.2f}ms | "
            f"Full p95: {interim_metrics.get('p95_full_latency_ms', 0):.2f}ms | "
            f"Chat: {interim_metrics.get('chat_acc', 0):.1f}% | "
            f"Task: {interim_metrics.get('task_acc', 0):.1f}% | "
            f"Dev: {interim_metrics.get('dev_acc', 0):.1f}% | "
            f"Doc: {interim_metrics.get('doc_acc', 0):.1f}%"
        )

        if stopped_by_token_cap:
            break

        if batch_index < len(batches):
            await asyncio.sleep(BATCH_GAP_SEC)

    subtest_elapsed_sec = time.perf_counter() - subtest_start
    progress.finish()
    return results, stopped_by_token_cap, subtest_elapsed_sec


async def run_difficulty(
    difficulty: str,
    client: httpx.AsyncClient,
    limiter: SlidingWindowRateLimiter
) -> Dict[str, Any]:
    requests_by_concurrency = build_requests_for_difficulty(difficulty)
    token_guard = TokenBudgetGuard(TOKEN_CAP_PER_DIFFICULTY_RUN)

    overall_results: List[EvalResult] = []
    stopped_early = False
    subtest_summaries: List[Dict[str, Any]] = []

    for concurrency in CONCURRENCY_LEVELS:
        sub_results, token_stop, sub_elapsed_sec = await run_subtest(
            difficulty=difficulty,
            concurrency=concurrency,
            requests=requests_by_concurrency[concurrency],
            client=client,
            limiter=limiter,
            token_guard=token_guard,
        )
        overall_results.extend(sub_results)

        metrics = compute_metrics(sub_results)
        subtest_summaries.append({
            "difficulty": difficulty,
            "concurrency": concurrency,
            "requests_completed": len(sub_results),
            "total_tokens": metrics.get("total_tokens", 0),
            "success_rate_pct": metrics.get("success_rate_pct", 0.0),
            "routing_accuracy_pct": metrics.get("routing_accuracy_pct", 0.0),
            "format_compliance_pct": metrics.get("format_compliance_pct", 0.0),
            "ttft_lt_500_pct": metrics.get("ttft_lt_500_pct", 0.0),
            "state_retention_pct": metrics.get("state_retention_pct", "N/A"),
            "p50_ttft_ms": metrics.get("p50_ttft_ms", 0),
            "p50_full_latency_ms": metrics.get("p50_full_latency_ms", 0),
            "p95_full_latency_ms": metrics.get("p95_full_latency_ms", 0),
            "subtest_elapsed_sec": sub_elapsed_sec,
            "chat_acc": metrics.get("chat_acc", 0.0),
            "task_acc": metrics.get("task_acc", 0.0),
            "dev_acc": metrics.get("dev_acc", 0.0),
            "doc_acc": metrics.get("doc_acc", 0.0),
            "stopped_by_token_cap": token_stop,
        })

        if token_stop:
            stopped_early = True
            print(f"Token cap reached during {difficulty} at concurrency {concurrency}. Remaining subtests will be skipped.")
            break

    difficulty_metrics = compute_metrics(overall_results)

    return {
        "difficulty": difficulty,
        "metrics": difficulty_metrics,
        "subtest_summary": subtest_summaries,
        "stopped_early": stopped_early,
        "requests_completed": len(overall_results),
        "tokens_consumed": token_guard.total_tokens,
        "results": overall_results,
    }


# =========================================================
# ENTRY POINT
# =========================================================
def choose_difficulties() -> List[str]:
    print("Select test scope:")
    print("1) Simple")
    print("2) Complex/Interrupt")
    print("3) Edge Case")
    print("4) All")
    choice = input("Enter choice [1-4]: ").strip()
    if choice == "1":
        return ["Simple"]
    if choice == "2":
        return ["Complex_Interrupt"]
    if choice == "3":
        return ["Edge_Case"]
    if choice == "4":
        return ["Simple", "Complex_Interrupt", "Edge_Case"]
    raise ValueError("Invalid choice. Please enter 1, 2, 3, or 4.")


async def async_main() -> None:
    if API_KEY == "PASTE_YOUR_CURRENT_API_KEY_HERE":
        print("\n[🚨 严重错误] 请先把你当前的 API_KEY 填回代码。")
        return

    selected = choose_difficulties()
    planned_requests = len(selected) * REQUESTS_PER_DIFFICULTY
    print(f"Planned requests for this run: {planned_requests}")
    print(f"Daily request budget assumption: {DAILY_REQUEST_BUDGET}")
    print(f"Token cap per difficulty run: {TOKEN_CAP_PER_DIFFICULTY_RUN}")

    if planned_requests > DAILY_REQUEST_BUDGET:
        print("WARNING: planned requests exceed the stated daily request budget.")
        answer = input("Continue anyway? [y/N]: ").strip().lower()
        if answer != "y":
            print("Aborted by user.")
            return

    run_root = ROOT_DIR / f"run_{timestamp_str()}"
    ensure_dir(run_root)
    limiter = SlidingWindowRateLimiter(GLOBAL_RPM)

    limits = httpx.Limits(
        max_keepalive_connections=HTTP_MAX_KEEPALIVE_CONNECTIONS,
        max_connections=HTTP_MAX_CONNECTIONS,
        keepalive_expiry=HTTP2_KEEPALIVE_EXPIRY_SEC,
    )
    timeout = httpx.Timeout(
        connect=10.0,
        read=REQUEST_TIMEOUT_SEC,
        write=10.0,
        pool=10.0,
    )

    difficulty_reports: List[Dict[str, Any]] = []
    all_results: List[EvalResult] = []

    async with httpx.AsyncClient(
        http2=True,
        limits=limits,
        timeout=timeout,
        headers={"Content-Type": "application/json"},
    ) as client:
        for difficulty in selected:
            print(f"\nStarting difficulty: {difficulty}")
            rep = await run_difficulty(difficulty, client, limiter)
            difficulty_reports.append(rep)
            all_results.extend(rep["results"])

    run_config = {
        "model": MODEL_NAME,
        "thinking_level": THINKING_LEVEL,
        "max_output_tokens": MAX_OUTPUT_TOKENS,
        "temperature": TEMPERATURE,
        "top_p": TOP_P,
        "max_history_messages": MAX_HISTORY_MESSAGES,
        "http2_enabled": True,
        "keepalive_expiry_sec": HTTP2_KEEPALIVE_EXPIRY_SEC,
        "streaming_enabled": True,
        "difficulties_tested": ", ".join(selected),
        "total_requests_evaluated": len(all_results),
        "requests_per_subtest": REQUESTS_PER_SUBTEST,
        "concurrency_levels": str(CONCURRENCY_LEVELS),
        "batch_gap_after_each_completed_batch_(sec)": BATCH_GAP_SEC,
        "global_RPM_limiter": GLOBAL_RPM,
        "request_timeout_(sec)": REQUEST_TIMEOUT_SEC,
        "max_retries": MAX_RETRIES,
        "daily_request_budget_assumption": DAILY_REQUEST_BUDGET,
        "token_cap_per_difficulty_run": TOKEN_CAP_PER_DIFFICULTY_RUN,
        "random_seed": SEED,
    }

    report_path = run_root / "Consolidated_Benchmark_Report.docx"
    generate_comprehensive_word_report(
        report_path,
        run_config,
        all_results,
        difficulty_reports,
    )

    print(f"\nRun completed. Output saved to: {run_root}")
    print("Generated file:")
    print(f" - {report_path}")


def main() -> None:
    asyncio.run(async_main())


if __name__ == "__main__":
    main()