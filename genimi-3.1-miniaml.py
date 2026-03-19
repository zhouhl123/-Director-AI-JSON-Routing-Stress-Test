import asyncio
import json
import math
import random
import time
from collections import Counter, defaultdict
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import sys

# 修复 Windows 下 aiohttp 退出时 Event loop is closed 的报错
if sys.platform == 'win32':
    asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())

import aiohttp
from docx import Document

# =========================================================
# LOCAL SETTINGS
# =========================================================
API_KEY = "AIzaSyCcHjM4w1DPJNPiEsejPtdj1pbrsMWU9yQ"  # 请确保填入你真实的 API KEY
MODEL_NAME = "gemini-3.1-flash-lite-preview"
ROOT_DIR = Path(r"C:\Users\Administrator\Desktop\AO AI\C\3")

# 【数值对齐】：36 * 3 = 108
REQUESTS_PER_DIFFICULTY = 108
REQUESTS_PER_SUBTEST = 36
CONCURRENCY_LEVELS = [1, 2, 3]
BATCH_GAP_SEC = 5

DAILY_REQUEST_BUDGET = 500
TOKEN_CAP_PER_DIFFICULTY_RUN = 250_000

GLOBAL_RPM = 10
MAX_RETRIES = 3
REQUEST_TIMEOUT_SEC = 60
SEED = 20260316

# 【有算力限制配置】：强制开启 minimal 模式，并把 Token 物理压缩到 100，逼迫模型“条件反射”
THINKING_LEVEL = "minimal"
MAX_OUTPUT_TOKENS = 100
TEMPERATURE = 0.0
TOP_P = 0.1

COMPLEXITY_POWER = 0

GEMINI_API_URL = (
    f"https://generativelanguage.googleapis.com/v1beta/models/"
    f"{MODEL_NAME}:generateContent?key={API_KEY}"
)

# 路线标签：彻底将 code 升级为 dev (Developer/Technical)
ROUTE_LABELS = ["chat", "task", "dev", "doc"]
DIFFICULTY_KEYS = ["Simple", "Complex_Interrupt", "Edge_Case"]

# =========================================================
# ADVERSARIAL SYSTEM PROMPT (对抗性陷阱)
# =========================================================
# 故意设置绝对的启发式死板规则，算力受限时，模型必然会被这些规则带偏
DIRECTOR_SYSTEM = """
SYSTEM ROLE: ROUTER.
DEFINITIONS:
- chat: casual conversation, greetings, or non-technical general knowledge.
- task: project management, planning, roadmaps, or non-technical actionable checklists.
- dev: software development, programming, debugging, architecture concepts, databases, and dev tools.
- doc: reading, summarizing, analyzing, or extracting from user-provided text.

ROUTING HEURISTIC & TRAPS:
- CRITICAL RULE 1: If the user's request explicitly mentions coding languages (e.g., Python, C++), software tools (e.g., Kubernetes, Docker), or APIs in ANY way, you MUST route to 'dev' IMMEDIATELY.
- CRITICAL RULE 2: If the request explicitly asks to build a "roadmap", "sprint plan", or "timeline", route to 'task' IMMEDIATELY.

OUTPUT RULES:
Return RAW JSON only. DO NOT wrap in ```json.
The 'reason' field MUST be a helpful 10-15 word direct answer or actionable summary.

EXAMPLE:
{"route":"chat","confidence":0.99,"needs_memory":false,"needs_tool":false,"reason":"Hello! I am ready to help you with your tasks today."}
""".strip()

random.seed(SEED)


# =========================================================
# DATA MODELS
# =========================================================
@dataclass
class EvalRequest:
    sample_id: str
    difficulty: str
    concurrency: int
    input_text: str
    expected_route: str
    conversation_id: str = ""
    turn_index: int = 1
    state_case: str = "none"
    history: Optional[List[Dict[str, str]]] = None


@dataclass
class EvalResult:
    sample_id: str
    difficulty: str
    concurrency: int
    conversation_id: str
    turn_index: int
    state_case: str
    input_text: str
    expected_route: str
    predicted_route: str
    parse_success: bool
    schema_valid: bool
    format_compliant: bool
    route_correct: bool
    state_correct: bool
    carry_correct: bool
    interrupt_correct: bool
    latency_ms: float
    latency_lt_500: bool
    attempts: int
    retry_count: int
    retry_success: bool
    success: bool
    rate_limited: bool
    error_type: str
    error_message: str
    raw_response: str
    total_tokens: int
    prompt_tokens: int
    candidate_tokens: int
    token_estimated: bool


# =========================================================
# INFRASTRUCTURE
# =========================================================
class SlidingWindowRateLimiter:
    def __init__(self, rpm: int):
        self.rpm = max(1, rpm)
        self.timestamps: List[float] = []
        self.lock = asyncio.Lock()

    async def acquire(self) -> None:
        while True:
            async with self.lock:
                now = time.monotonic()
                # Remove timestamps older than 60 seconds
                self.timestamps = [t for t in self.timestamps if now - t < 60.0]
                if len(self.timestamps) < self.rpm:
                    self.timestamps.append(now)
                    return

                # If we've hit the limit, calculate how long to wait
                wait_seconds = 60.0 - (now - self.timestamps[0]) + 0.01

            # Wait outside the lock so other tasks can proceed
            await asyncio.sleep(max(wait_seconds, 0.05))


class TokenBudgetGuard:
    def __init__(self, token_cap: int):
        self.token_cap = max(1, token_cap)
        self.total_tokens = 0
        self.lock = asyncio.Lock()

    async def add_and_check(self, tokens: int) -> bool:
        async with self.lock:
            self.total_tokens += max(0, tokens)
            return self.total_tokens <= self.token_cap


def timestamp_str() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def clean_json_response(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        lines = text.splitlines()
        if lines and lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].startswith("```"):
            lines = lines[:-1]
        text = "\n".join(lines).strip()
    return text


def quantile(data: List[float], q: float) -> float:
    if not data:
        return 0.0
    s = sorted(data)
    pos = (len(s) - 1) * q
    lo = int(math.floor(pos))
    hi = int(math.ceil(pos))
    if lo == hi:
        return s[lo]
    return s[lo] + (s[hi] - s[lo]) * (pos - lo)


def chunked(items: List[Any], size: int) -> List[List[Any]]:
    return [items[i:i + size] for i in range(0, len(items), size)]


def route_index(route: str) -> int:
    try:
        return ROUTE_LABELS.index(route)
    except ValueError:
        return -1


def rough_token_estimate(text: str) -> int:
    if not text:
        return 0
    return max(1, math.ceil(len(text) / 4))


def estimate_usage_fallback(request: EvalRequest, response_text: str) -> Tuple[int, int, int]:
    history_text = json.dumps(request.history or [], ensure_ascii=False)
    prompt_text = DIRECTOR_SYSTEM + "\n" + history_text + "\n" + request.input_text
    prompt_tokens = rough_token_estimate(prompt_text)
    candidate_tokens = rough_token_estimate(response_text)
    total_tokens = prompt_tokens + candidate_tokens
    return total_tokens, prompt_tokens, candidate_tokens


class ProgressPrinter:
    def __init__(self, label: str, total: int):
        self.label = label
        self.total = total
        self.current = 0

    def update(self, n: int = 1, suffix: str = "") -> None:
        self.current += n
        pct = (self.current / self.total * 100.0) if self.total else 100.0
        width = 30
        filled = int(width * self.current / self.total) if self.total else width
        bar = "[" + "#" * filled + "-" * (width - filled) + "]"
        msg = f"\r{self.label} {bar} {self.current}/{self.total} ({pct:.1f}%) {suffix}"
        print(msg, end="", flush=True)

    def finish(self) -> None:
        print()


# =========================================================
# PROMPT BUILDING
# =========================================================
def cube_wrap(base_text: str, route: str, difficulty: str, variant: int, extra_block: str = "") -> str:
    combined = base_text.strip()
    if extra_block:
        combined += "\n" + extra_block.strip()
    return combined


# =========================================================
# DATASET GENERATION (ADVERSARIAL REDESIGN)
# =========================================================
def build_simple_base() -> List[EvalRequest]:
    items: List[EvalRequest] = []

    chat_topics = [
                      "say hello", "introduce yourself", "give a short greeting", "answer casually",
                      "tell me a fun fact", "how are you", "say good morning", "reply politely", "chat me up"
                  ] * 4

    task_topics = [
                      "make a 2-week sprint plan", "create a marketing roadmap", "break this task into steps",
                      "make a checklist for event planning", "plan a research workflow", "outline a presentation",
                      "draft an implementation plan", "create a weekly execution plan", "make a launch checklist"
                  ] * 4

    dev_topics = [
                     "Write a Python script to parse csv", "Explain REST vs GraphQL", "Debug this React component",
                     "How does Kubernetes handle pod scaling?", "Write a bash script", "What is MVC architecture?",
                     "Build a Node API", "Explain JavaScript closures", "Fix this C++ memory leak"
                 ] * 4

    doc_topics = [
                     "summarize these meeting notes", "review this memo and extract actions",
                     "analyze this financial document",
                     "extract the key facts from this note", "summarize the following text",
                     "review this quarterly report",
                     "extract decisions from this document", "summarize the meeting content",
                     "give me an abstract of this article"
                 ] * 4

    for i in range(36):
        items.append(EvalRequest(f"simple_chat_{i + 1:03d}", "Simple", 0, chat_topics[i], "chat"))
        items.append(EvalRequest(f"simple_task_{i + 1:03d}", "Simple", 0, task_topics[i], "task"))
        items.append(EvalRequest(f"simple_dev_{i + 1:03d}", "Simple", 0, dev_topics[i], "dev"))
        items.append(EvalRequest(f"simple_doc_{i + 1:03d}", "Simple", 0, doc_topics[i], "doc"))

    return items


def _conversation_to_turn_requests(
        difficulty: str,
        templates: List[List[Dict[str, Any]]],
        request_count: int
) -> List[EvalRequest]:
    base_turns: List[EvalRequest] = []

    for conv_idx, conv in enumerate(templates, start=1):
        history: List[Dict[str, str]] = []
        conversation_id = f"{difficulty.lower()}_conv_{conv_idx:03d}"

        for turn_idx, turn in enumerate(conv, start=1):
            base_turns.append(EvalRequest(
                sample_id=f"{conversation_id}_turn_{turn_idx}",
                difficulty=difficulty,
                concurrency=0,
                input_text=turn["text"],
                expected_route=turn["expected_route"],
                conversation_id=conversation_id,
                turn_index=turn_idx,
                state_case=turn.get("state_case", "none"),
                history=[dict(x) for x in history],
            ))
            history.append({"role": "user", "text": turn["text"]})
            history.append({"role": "model", "text": "Acknowledged."})

    items: List[EvalRequest] = []
    for i in range(request_count):
        src = base_turns[i % len(base_turns)]
        items.append(EvalRequest(
            sample_id=f"{src.sample_id}_rep_{i + 1:03d}",
            difficulty=src.difficulty,
            concurrency=0,
            input_text=src.input_text,
            expected_route=src.expected_route,
            conversation_id=src.conversation_id,
            turn_index=src.turn_index,
            state_case=src.state_case,
            history=[dict(x) for x in (src.history or [])],
        ))
    return items


def build_complex_templates() -> List[List[Dict[str, Any]]]:
    templates: List[List[Dict[str, Any]]] = []

    bases = {
        "chat": "Hello! I am feeling a bit tired today. Can we just chat?",
        "task": "Create a 2-week launch plan for the new retail store.",
        "dev": "Explain the difference between REST and GraphQL APIs.",
        "doc": "Summarize the attached HR meeting notes into 4 bullet points."
    }
    carries = {
        "chat": "Actually, reply to me as if you are a grumpy pirate.",
        "task": "Shorten the timeline to 1 week but keep the same format.",
        "dev": "Now compare them regarding caching capabilities and speed.",
        "doc": "Make the summary shorter and targeted for the CEO."
    }
    interrupts = {
        "chat": "Stop acting like a pirate. Tell me a completely different joke.",
        "task": "Forget the launch plan. Break down a marketing event into a checklist.",
        "dev": "Wait, ignore the APIs. Write a SQL query to fetch user data.",
        "doc": "Actually switch tasks. Read this policy writeup and extract the dates."
    }

    intents = ["chat", "task", "dev", "doc"]

    for i in range(12):
        base_intent = intents[i % 4]
        interrupt_intent = intents[(i + 1) % 4]

        templates.append([
            {"text": cube_wrap(bases[base_intent], base_intent, "Complex_Interrupt", i),
             "expected_route": base_intent, "state_case": "none"},
            {"text": cube_wrap(carries[base_intent], base_intent, "Complex_Interrupt", i + 100),
             "expected_route": base_intent, "state_case": "carry"},
            {"text": cube_wrap(interrupts[interrupt_intent], interrupt_intent, "Complex_Interrupt", i + 200),
             "expected_route": interrupt_intent, "state_case": "interrupt"},
        ])

    return templates


def build_edge_templates() -> List[List[Dict[str, Any]]]:
    templates: List[List[Dict[str, Any]]] = []

    # 终极对抗模式：包含系统规则 1 和 2 的硬性触发词，但实际意图完全相反
    bases = {
        "chat": [
            "I know I asked you to write a Python script and build a Jira roadmap yesterday, but honestly I'm just depressed today. Can we ignore the code and just talk like friends?",
            "Forget about the Docker Kubernetes pod configurations. I'm taking a break. How is your day going as an AI?"
        ],
        "task": [
            "I need you to read this 50-page C++ documentation and the PostgreSQL schema. Then, entirely ignore the code and give me a 4-week project management timeline for hiring the team.",
            "My boss wants a Python backend built. Provide a milestone checklist of the project phases. Do not write any code, strictly project planning."
        ],
        "dev": [
            "Hey buddy! Long time no see. Hope you are having a wonderful morning. Listen, I know we usually chat casually, but could you quickly output a Bash script to recursively delete logs?",
            "I want to make a roadmap for learning. Step 1 is to understand closures. Just execute Step 1 immediately by giving me a JavaScript code example."
        ],
        "doc": [
            "Create a timeline roadmap. Step 1: Read the attached legal contract. Step 2: Extract the termination clauses. Ignore the roadmap planning and just execute Step 2 on this text.",
            "Write a C++ script that parses JSON... wait, nevermind. Just read this error log text block and summarize the warnings into 3 bullet points."
        ]
    }
    carries = {
        "chat": [
            "Yeah exactly, keep the casual vibe going. Just tell me a random non-technical story.",
            "Thanks for listening. Tell me a joke to cheer me up, no tech talk."
        ],
        "task": [
            "Shorten that team hiring timeline to 2 weeks, but again, no technical concepts.",
            "Add a QA management phase to that checklist. Still strictly planning."
        ],
        "dev": [
            "Awesome, now add robust try-catch error handling to that Bash script.",
            "Now optimize that JS code snippet you just provided."
        ],
        "doc": [
            "Re-do the extraction from that contract but group them by monetary penalties.",
            "Shorten the summary of the log file you just read to one sentence."
        ]
    }
    interrupts = {
        "chat": [
            "Okay, break is over. Write a fully functional React login component.",
            "Stop chatting. Create a 3-month product sprint roadmap for my startup."
        ],
        "task": [
            "Wait, forget the hiring roadmap. Let's just have a friendly, casual chat about art.",
            "Pause the checklist. Read this Python error transcript and summarize the main stack trace issue."
        ],
        "dev": [
            "Stop coding. Create a Gantt chart timeline for deploying the servers.",
            "Forget the Bash script. I just want to vent. Can we talk about my stress?"
        ],
        "doc": [
            "Stop reading the logs. Write a SQL query to delete all users.",
            "Halt the summary. Make a 14-day execution plan for fixing the issues."
        ]
    }

    intents = ["chat", "task", "dev", "doc"]

    for i in range(12):
        base_intent = intents[i % 4]
        interrupt_intent = intents[(i + 2) % 4]
        b_idx = i % 2

        templates.append([
            {"text": cube_wrap(bases[base_intent][b_idx], base_intent, "Edge_Case", i),
             "expected_route": base_intent, "state_case": "none"},
            {"text": cube_wrap(carries[base_intent][b_idx], base_intent, "Edge_Case", i + 100),
             "expected_route": base_intent, "state_case": "carry"},
            {"text": cube_wrap(interrupts[interrupt_intent][b_idx], interrupt_intent, "Edge_Case", i + 200),
             "expected_route": interrupt_intent, "state_case": "interrupt"},
        ])

    return templates


def build_requests_for_difficulty(difficulty: str) -> Dict[int, List[EvalRequest]]:
    def simple_subtest(c: int) -> List[EvalRequest]:
        base = build_simple_base()
        out: List[EvalRequest] = []
        for i in range(REQUESTS_PER_SUBTEST):
            src = base[i % len(base)]
            text = src.input_text + f"\n[Benchmark Tag: simple-subtest-c{c}-variant-{i}]"
            out.append(EvalRequest(
                sample_id=f"{src.sample_id}_c{c}_rep_{i + 1:03d}",
                difficulty="Simple",
                concurrency=c,
                input_text=text,
                expected_route=src.expected_route,
            ))
        return out

    def complex_subtest(c: int) -> List[EvalRequest]:
        out = _conversation_to_turn_requests("Complex_Interrupt", build_complex_templates(), REQUESTS_PER_SUBTEST)
        for item in out:
            item.concurrency = c
            item.sample_id = f"{item.sample_id}_c{c}"
        return out

    def edge_subtest(c: int) -> List[EvalRequest]:
        out = _conversation_to_turn_requests("Edge_Case", build_edge_templates(), REQUESTS_PER_SUBTEST)
        for item in out:
            item.concurrency = c
            item.sample_id = f"{item.sample_id}_c{c}"
        return out

    if difficulty == "Simple":
        return {c: simple_subtest(c) for c in CONCURRENCY_LEVELS}
    if difficulty == "Complex_Interrupt":
        return {c: complex_subtest(c) for c in CONCURRENCY_LEVELS}
    if difficulty == "Edge_Case":
        return {c: edge_subtest(c) for c in CONCURRENCY_LEVELS}

    raise ValueError(f"Unsupported difficulty: {difficulty}")


# =========================================================
# GEMINI DIRECTOR CALL
# =========================================================
def build_contents(history: Optional[List[Dict[str, str]]], user_text: str) -> List[Dict[str, Any]]:
    contents: List[Dict[str, Any]] = []

    for turn in (history or [])[-8:]:
        contents.append({"role": turn["role"], "parts": [{"text": turn["text"]}]})

    wrapped_text = f"Classify this user intent:\n<User_Message>\n{user_text.strip()}\n</User_Message>"
    contents.append({"role": "user", "parts": [{"text": wrapped_text}]})
    return contents


def extract_text_from_gemini(data: Dict[str, Any]) -> str:
    try:
        parts = data["candidates"][0]["content"]["parts"]
        return "".join([p.get("text", "") for p in parts if "text" in p]).strip()
    except Exception:
        return ""


def extract_usage_metadata(data: Dict[str, Any]) -> Tuple[int, int, int]:
    md = data.get("usageMetadata", {}) or {}
    total_tokens = int(md.get("totalTokenCount", 0) or 0)
    prompt_tokens = int(md.get("promptTokenCount", 0) or 0)
    candidate_tokens = int(md.get("candidatesTokenCount", 0) or 0)
    return total_tokens, prompt_tokens, candidate_tokens


def validate_schema(obj: Dict[str, Any]) -> bool:
    required = ["route", "confidence", "needs_memory", "needs_tool", "reason"]
    if not isinstance(obj, dict):
        return False
    if not all(k in obj for k in required):
        return False
    if obj.get("route") not in ROUTE_LABELS:
        return False
    return True


async def call_director_once(
        session: aiohttp.ClientSession,
        limiter: SlidingWindowRateLimiter,
        request: EvalRequest
) -> Tuple[int, str, float]:
    await limiter.acquire()

    # 【重启封印】：加入 thinkingConfig 并设置 minimal，强制剥夺逻辑推演空间
    payload: Dict[str, Any] = {
        "systemInstruction": {"parts": [{"text": DIRECTOR_SYSTEM}]},
        "contents": build_contents(request.history, request.input_text),
        "generationConfig": {
            "temperature": TEMPERATURE,
            "topP": TOP_P,
            "maxOutputTokens": MAX_OUTPUT_TOKENS,
            "responseMimeType": "application/json",
            "thinkingConfig": {
                "thinkingLevel": THINKING_LEVEL
            }
        },
    }

    start_t = time.perf_counter()
    try:
        async with session.post(
                GEMINI_API_URL,
                headers={"Content-Type": "application/json"},
                json=payload,
                timeout=aiohttp.ClientTimeout(total=REQUEST_TIMEOUT_SEC),
        ) as resp:
            text = await resp.text()
            latency = (time.perf_counter() - start_t) * 1000.0
            return resp.status, text, latency
    except Exception as e:
        latency = (time.perf_counter() - start_t) * 1000.0
        raise e


async def evaluate_request(
        session: aiohttp.ClientSession,
        limiter: SlidingWindowRateLimiter,
        request: EvalRequest
) -> EvalResult:
    attempts = 0
    retry_count = 0
    rate_limited = False
    raw_response = ""
    error_type = ""
    error_message = ""
    total_tokens = 0
    prompt_tokens = 0
    candidate_tokens = 0
    token_estimated = False
    predicted_route = ""
    parse_success = False
    schema_valid = False
    pure_latency_ms = 0.0

    while attempts < 1 + MAX_RETRIES:
        attempts += 1
        try:
            status, body, pure_latency_ms = await call_director_once(session, limiter, request)

            if status != 200:
                print(f"\n[🚨 API ERROR] HTTP {status}: {body[:200]}")

            raw_response = body

            if status == 429 or "RESOURCE_EXHAUSTED" in body or "rate limit" in body.lower():
                rate_limited = True
                error_type = "rate_limit"
                error_message = body[:1000]
                if attempts <= MAX_RETRIES:
                    retry_count += 1
                    await asyncio.sleep(min(2 ** (attempts - 1), 10))
                    continue
            elif status >= 500:
                error_type = f"http_{status}"
                error_message = body[:1000]
                if attempts <= MAX_RETRIES:
                    retry_count += 1
                    await asyncio.sleep(min(2 ** (attempts - 1), 10))
                    continue
            elif status != 200:
                error_type = f"http_{status}"
                error_message = body[:1000]
                break

            data = json.loads(body)
            text = extract_text_from_gemini(data)
            clean_text = clean_json_response(text)

            total_tokens, prompt_tokens, candidate_tokens = extract_usage_metadata(data)
            if total_tokens <= 0:
                total_tokens, prompt_tokens, candidate_tokens = estimate_usage_fallback(request, clean_text or body)
                token_estimated = True

            try:
                obj = json.loads(clean_text)
                parse_success = True
                schema_valid = validate_schema(obj)

                if isinstance(obj, dict):
                    predicted_route = str(obj.get("route", ""))
                else:
                    predicted_route = ""

                if not schema_valid:
                    error_type = "schema_invalid"
                    error_message = clean_text[:1000]
                else:
                    error_type = ""
                    error_message = ""
                raw_response = clean_text
            except json.JSONDecodeError as e:
                error_type = "json_parse_error"
                error_message = f"{e}: {clean_text[:1000]}"
                break
            break

        except asyncio.TimeoutError:
            error_type = "timeout"
            error_message = f"Timeout after {REQUEST_TIMEOUT_SEC} seconds"
            if attempts <= MAX_RETRIES:
                retry_count += 1
                await asyncio.sleep(min(2 ** (attempts - 1), 10))
                continue
            break

        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)[:1000]
            print(f"\n[🚨 CODE ERROR] {error_message}")
            if attempts <= MAX_RETRIES:
                retry_count += 1
                await asyncio.sleep(min(2 ** (attempts - 1), 10))
                continue
            break

    format_compliant = parse_success and schema_valid
    route_correct = format_compliant and predicted_route == request.expected_route

    if request.state_case == "carry":
        carry_correct = route_correct
        interrupt_correct = False
        state_correct = carry_correct
    elif request.state_case == "interrupt":
        carry_correct = False
        interrupt_correct = route_correct
        state_correct = interrupt_correct
    else:
        carry_correct = False
        interrupt_correct = False
        state_correct = False

    success = format_compliant and error_type == ""
    retry_success = success and retry_count > 0

    return EvalResult(
        sample_id=request.sample_id,
        difficulty=request.difficulty,
        concurrency=request.concurrency,
        conversation_id=request.conversation_id,
        turn_index=request.turn_index,
        state_case=request.state_case,
        input_text=request.input_text,
        expected_route=request.expected_route,
        predicted_route=predicted_route,
        parse_success=parse_success,
        schema_valid=schema_valid,
        format_compliant=format_compliant,
        route_correct=route_correct,
        state_correct=state_correct,
        carry_correct=carry_correct,
        interrupt_correct=interrupt_correct,
        latency_ms=pure_latency_ms,
        latency_lt_500=(0 < pure_latency_ms < 500.0),
        attempts=attempts,
        retry_count=retry_count,
        retry_success=retry_success,
        success=success,
        rate_limited=rate_limited,
        error_type=error_type,
        error_message=error_message,
        raw_response=raw_response,
        total_tokens=total_tokens,
        prompt_tokens=prompt_tokens,
        candidate_tokens=candidate_tokens,
        token_estimated=token_estimated,
    )


# =========================================================
# METRICS & REPORTING LOGIC
# =========================================================
def compute_metrics(results: List[EvalResult]) -> Dict[str, Any]:
    n = len(results)
    if n == 0:
        return {}

    fmt = sum(r.format_compliant for r in results)
    parse_ok = sum(r.parse_success for r in results)
    schema_ok = sum(r.schema_valid for r in results)
    route_ok = sum(r.route_correct for r in results)
    valid_only_den = max(1, schema_ok)

    state_results = [r for r in results if r.state_case != "none"]
    carry_results = [r for r in results if r.state_case == "carry"]
    interrupt_results = [r for r in results if r.state_case == "interrupt"]

    sr_pct = round(sum(r.state_correct for r in state_results) / len(state_results) * 100.0,
                   2) if state_results else "N/A"
    co_pct = round(sum(r.carry_correct for r in carry_results) / len(carry_results) * 100.0,
                   2) if carry_results else "N/A"
    in_pct = round(sum(r.interrupt_correct for r in interrupt_results) / len(interrupt_results) * 100.0,
                   2) if interrupt_results else "N/A"

    latencies = [r.latency_ms for r in results if r.success and r.latency_ms > 0]

    success_count = sum(r.success for r in results)
    rate_limit_hits = sum(r.rate_limited for r in results)
    retry_success_count = sum(r.retry_success for r in results)
    total_retries = sum(r.retry_count for r in results)
    tokens_total = sum(r.total_tokens for r in results)
    prompt_tokens_total = sum(r.prompt_tokens for r in results)
    candidate_tokens_total = sum(r.candidate_tokens for r in results)
    estimated_token_rows = sum(r.token_estimated for r in results)

    confusion = [[0 for _ in ROUTE_LABELS] for _ in ROUTE_LABELS]
    for r in results:
        ei = route_index(r.expected_route)
        pi = route_index(r.predicted_route)
        if ei >= 0 and pi >= 0:
            confusion[ei][pi] += 1

    cat_acc = {}
    for route in ROUTE_LABELS:
        cat_results = [r for r in results if r.expected_route == route]
        if not cat_results:
            cat_acc[route] = 0.0
        else:
            cat_acc[route] = round(sum(r.route_correct for r in cat_results) / len(cat_results) * 100.0, 2)

    return {
        "total_requests": n,
        "format_compliance_pct": round(fmt / n * 100.0, 2),
        "parse_success_pct": round(parse_ok / n * 100.0, 2),
        "schema_pass_pct": round(schema_ok / n * 100.0, 2),
        "routing_accuracy_pct": round(route_ok / n * 100.0, 2),
        "valid_only_routing_accuracy_pct": round(route_ok / valid_only_den * 100.0, 2),
        "state_retention_pct": sr_pct,
        "carry_over_correctness_pct": co_pct,
        "interrupt_correctness_pct": in_pct,
        "latency_lt_500_pct": round(
            sum(r.latency_lt_500 for r in results if r.success) / max(1, len(latencies)) * 100.0,
            2) if latencies else 0.0,
        "mean_latency_ms": round(sum(latencies) / len(latencies), 2) if latencies else 0.0,
        "p50_latency_ms": round(quantile(latencies, 0.50), 2) if latencies else 0.0,
        "p95_latency_ms": round(quantile(latencies, 0.95), 2) if latencies else 0.0,
        "success_rate_pct": round(success_count / n * 100.0, 2),
        "rate_limit_hit_rate_pct": round(rate_limit_hits / n * 100.0, 2),
        "retry_success_rate_pct": round(
            retry_success_count / max(1, sum(r.retry_count > 0 for r in results)) * 100.0, 2
        ),
        "avg_retry_count": round(total_retries / n, 4),
        "total_tokens": tokens_total,
        "total_prompt_tokens": prompt_tokens_total,
        "total_candidate_tokens": candidate_tokens_total,
        "avg_tokens_per_request": round(tokens_total / n, 2),
        "estimated_token_rows": estimated_token_rows,
        "confusion_matrix": confusion,
        "error_counts": dict(Counter(r.error_type or "none" for r in results)),
        "chat_acc": cat_acc.get("chat", 0.0),
        "task_acc": cat_acc.get("task", 0.0),
        "dev_acc": cat_acc.get("dev", 0.0),
        "doc_acc": cat_acc.get("doc", 0.0),
    }


def build_analysis(metrics: Dict[str, Any], difficulty: str, concurrency: Optional[int] = None) -> List[str]:
    lines: List[str] = []
    prefix = f"{difficulty}"
    if concurrency is not None:
        prefix += f" @ concurrency {concurrency}"

    lines.append(f"This section summarizes the routing benchmark for {prefix}.")
    lines.append(
        f"Format Compliance was {metrics['format_compliance_pct']}%, while Routing Accuracy was "
        f"{metrics['routing_accuracy_pct']}%. Valid-Only Routing Accuracy was "
        f"{metrics['valid_only_routing_accuracy_pct']}%."
    )
    lines.append(
        f"API Latency below 500 ms was achieved in {metrics['latency_lt_500_pct']}% of successful requests, "
        f"with p50 = {metrics['p50_latency_ms']} ms and p95 = {metrics['p95_latency_ms']} ms."
    )

    state_ret = metrics['state_retention_pct']
    carry_ret = metrics['carry_over_correctness_pct']
    int_ret = metrics['interrupt_correctness_pct']

    if state_ret != "N/A":
        lines.append(
            f"State Retention was {state_ret}%, with Carry-over Correctness = "
            f"{carry_ret}% and Interrupt Correctness = {int_ret}%."
        )
    else:
        lines.append("State Retention metrics are N/A (no multi-turn requests in this subset).")

    lines.append(
        f"Success Rate was {metrics['success_rate_pct']}%, Rate-Limit Hit Rate was "
        f"{metrics['rate_limit_hit_rate_pct']}%, and the average retry count was "
        f"{metrics['avg_retry_count']}."
    )
    lines.append(
        f"Total tokens consumed were {metrics['total_tokens']}, averaging "
        f"{metrics['avg_tokens_per_request']} tokens per request."
    )
    return lines


# =========================================================
# WORD REPORTS ONLY
# =========================================================
def add_key_value_table(doc: Document, title: str, data: Dict[str, Any]) -> None:
    doc.add_heading(title, level=3)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    for k, v in data.items():
        if isinstance(v, float):
            val_str = f"{v:.2f}"
        else:
            val_str = str(v)
        row = table.add_row().cells
        row[0].text = str(k).replace("_", " ").title()
        row[1].text = val_str
    doc.add_paragraph()


def generate_comprehensive_word_report(
        out_path: Path,
        run_config: Dict[str, Any],
        all_results: List[EvalResult],
        difficulty_reports: List[Dict[str, Any]]
) -> None:
    overall_metrics = compute_metrics(all_results)
    analysis_lines = build_analysis(overall_metrics, "Overall")

    doc = Document()
    doc.add_heading("Adversarial Routing Benchmark Report (Compute Limited)", level=1)

    doc.add_heading("Test Configuration", level=2)
    for k, v in run_config.items():
        if k == "root_output": continue
        doc.add_paragraph(f"{k.replace('_', ' ').capitalize()}: {v}")

    doc.add_heading("Overall Core KPIs", level=2)
    core_kpis = {
        "Format Compliance (%)": overall_metrics.get('format_compliance_pct', 0),
        "Routing Accuracy (%)": overall_metrics.get('routing_accuracy_pct', 0),
        "State Retention (%)": overall_metrics.get('state_retention_pct', 'N/A'),
        "API Latency < 500 ms (%)": overall_metrics.get('latency_lt_500_pct', 0),
        "Total Tokens Consumed": overall_metrics.get('total_tokens', 0)
    }
    add_key_value_table(doc, "High-Level Overview", core_kpis)

    doc.add_heading("Overall Analysis", level=2)
    for line in analysis_lines:
        doc.add_paragraph(line)

    doc.add_page_break()

    doc.add_heading("Metrics Breakdown by Difficulty", level=2)
    for rep in difficulty_reports:
        m = rep["metrics"]
        diff_name = rep["difficulty"]
        doc.add_heading(f"Difficulty: {diff_name}", level=3)
        doc.add_paragraph(
            f"Format Compliance: {m.get('format_compliance_pct', 0):.2f}% | "
            f"Routing Accuracy: {m.get('routing_accuracy_pct', 0):.2f}% | "
            f"API Latency <500ms: {m.get('latency_lt_500_pct', 0):.2f}%"
        )

        doc.add_paragraph("Table 1: Execution & Resource Tracking").bold = True
        t1 = doc.add_table(rows=1, cols=3)
        t1.style = 'Table Grid'
        h1 = t1.rows[0].cells
        h1[0].text, h1[1].text, h1[2].text = "Concurrency", "Total Time (s)", "Total Tokens"
        for s in rep["subtest_summary"]:
            r1 = t1.add_row().cells
            r1[0].text = str(s["concurrency"])
            r1[1].text = f"{s.get('subtest_elapsed_sec', 0):.2f}"
            r1[2].text = str(s.get("total_tokens", 0))
        doc.add_paragraph()

        doc.add_paragraph("Table 2: Core Routing & Quality KPIs").bold = True
        t2 = doc.add_table(rows=1, cols=4)
        t2.style = 'Table Grid'
        h2 = t2.rows[0].cells
        h2[0].text, h2[1].text, h2[2].text, h2[
            3].text = "Concurrency", "Format Pass %", "Route Accuracy %", "State Retention %"
        for s in rep["subtest_summary"]:
            r2 = t2.add_row().cells
            r2[0].text = str(s["concurrency"])
            r2[1].text = f"{s.get('format_compliance_pct', 0):.2f}"
            r2[2].text = f"{s.get('routing_accuracy_pct', 0):.2f}"
            sr_val = s.get('state_retention_pct', 'N/A')
            r2[3].text = f"{sr_val:.2f}" if isinstance(sr_val, (int, float)) else str(sr_val)
        doc.add_paragraph()

        doc.add_paragraph("Table 3: Latency Distribution").bold = True
        t3 = doc.add_table(rows=1, cols=4)
        t3.style = 'Table Grid'
        h3 = t3.rows[0].cells
        h3[0].text, h3[1].text, h3[2].text, h3[3].text = "Concurrency", "< 500ms %", "p50 (ms)", "p95 (ms)"
        for s in rep["subtest_summary"]:
            r3 = t3.add_row().cells
            r3[0].text = str(s["concurrency"])
            r3[1].text = f"{s.get('latency_lt_500_pct', 0):.2f}"
            r3[2].text = f"{s.get('p50_latency_ms', 0):.2f}"
            r3[3].text = f"{s.get('p95_latency_ms', 0):.2f}"
        doc.add_paragraph()

        doc.add_paragraph("Table 4: Accuracy By Category").bold = True
        t4 = doc.add_table(rows=1, cols=5)
        t4.style = 'Table Grid'
        h4 = t4.rows[0].cells
        h4[0].text, h4[1].text, h4[2].text, h4[3].text, h4[
            4].text = "Concurrency", "Chat Acc %", "Task Acc %", "Dev Acc %", "Doc Acc %"
        for s in rep["subtest_summary"]:
            r4 = t4.add_row().cells
            r4[0].text = str(s["concurrency"])
            r4[1].text = f"{s.get('chat_acc', 0):.2f}"
            r4[2].text = f"{s.get('task_acc', 0):.2f}"
            r4[3].text = f"{s.get('dev_acc', 0):.2f}"
            r4[4].text = f"{s.get('doc_acc', 0):.2f}"

        doc.add_paragraph("-" * 40)

    doc.add_page_break()
    doc.add_heading("Appendix: Detailed Request KPIs & Transcripts", level=1)

    for i, r in enumerate(all_results, 1):
        doc.add_heading(f"Request #{i}: {r.sample_id} ({r.difficulty} - Concurrency {r.concurrency})", level=3)

        kpi_paragraph = doc.add_paragraph()

        kpi_paragraph.add_run("🎯 Core Match: ").bold = True
        kpi_paragraph.add_run(
            f"Expected [{r.expected_route}] vs Predicted [{r.predicted_route}]  =>  Correct: {'Yes' if r.route_correct else 'No'}\n")

        kpi_paragraph.add_run("⚡ Performance: ").bold = True
        kpi_paragraph.add_run(
            f"API Latency: {r.latency_ms:.2f} ms | Tokens: {r.total_tokens} (Prompt: {r.prompt_tokens}, Candidate: {r.candidate_tokens})\n")

        kpi_paragraph.add_run("🛠️ Quality Check: ").bold = True
        kpi_paragraph.add_run(
            f"Format Compliant: {'Yes' if r.format_compliant else 'No'} | JSON Parsed: {'Yes' if r.parse_success else 'No'} | Schema Valid: {'Yes' if r.schema_valid else 'No'}\n")

        if r.state_case != "none":
            kpi_paragraph.add_run("🔄 Multi-turn State: ").bold = True
            kpi_paragraph.add_run(
                f"Type: {r.state_case.capitalize()} | State Retention Correct: {'Yes' if r.state_correct else 'No'}\n")

        kpi_paragraph.add_run("🚨 Network & Errors: ").bold = True
        err_val = r.error_type if r.error_type else "None"
        kpi_paragraph.add_run(
            f"Attempts: {r.attempts} | Retries: {r.retry_count} | Rate Limited: {'Yes' if r.rate_limited else 'No'} | Error Type: {err_val}")

        doc.add_paragraph("[INPUT TEXT]").bold = True
        doc.add_paragraph(r.input_text or "N/A")

        doc.add_paragraph("[RAW OUTPUT]").bold = True
        doc.add_paragraph(r.raw_response or "N/A")

        doc.add_paragraph("-" * 60)

    doc.save(str(out_path))


# =========================================================
# RUNNERS
# =========================================================
async def run_subtest(
        difficulty: str,
        concurrency: int,
        requests: List[EvalRequest],
        limiter: SlidingWindowRateLimiter,
        token_guard: TokenBudgetGuard,
) -> Tuple[List[EvalResult], bool, float]:
    batches = chunked(requests, concurrency)
    progress = ProgressPrinter(f"{difficulty} @ concurrency {concurrency}", len(requests))
    results: List[EvalResult] = []
    stopped_by_token_cap = False

    subtest_start = time.perf_counter()

    connector = aiohttp.TCPConnector(limit=0, keepalive_timeout=300, enable_cleanup_closed=True)
    async with aiohttp.ClientSession(connector=connector) as session:
        for batch_index, batch in enumerate(batches, start=1):
            batch_results = await asyncio.gather(*(evaluate_request(session, limiter, req) for req in batch))
            batch_tokens = 0

            for r in batch_results:
                batch_tokens += r.total_tokens
                within_budget = await token_guard.add_and_check(r.total_tokens)
                results.append(r)

                # 实时打印被误判的请求，方便观察陷阱的命中率
                if r.success and not r.route_correct:
                    print(
                        f"\n[⚠️ 劫持成功] 期望: {r.expected_route.upper()} | 实际判为: {r.predicted_route.upper()} | 测试原句: '{r.input_text}'")

                progress.update(
                    1,
                    suffix=f"tokens={token_guard.total_tokens}/{token_guard.token_cap}"
                )
                if not within_budget:
                    stopped_by_token_cap = True

            interim_metrics = compute_metrics(results)
            sr_val = interim_metrics.get('state_retention_pct', 'N/A')
            sr_str = f"{sr_val:.1f}%" if isinstance(sr_val, (int, float)) else str(sr_val)
            print(
                f"\n  [Batch {batch_index}/{len(batches)}] Tokens: +{batch_tokens} (Total: {token_guard.total_tokens})\n"
                f"  => 实时 KPI: Format: {interim_metrics.get('format_compliance_pct', 0):.1f}% | "
                f"Route Acc: {interim_metrics.get('routing_accuracy_pct', 0):.1f}% | "
                f"State: {sr_str} | <500ms: {interim_metrics.get('latency_lt_500_pct', 0):.1f}% | "
                f"p50: {interim_metrics.get('p50_latency_ms', 0):.2f}ms | p95: {interim_metrics.get('p95_latency_ms', 0):.2f}ms | "
                f"Chat: {interim_metrics.get('chat_acc', 0):.1f}% | Task: {interim_metrics.get('task_acc', 0):.1f}% | "
                f"Dev: {interim_metrics.get('dev_acc', 0):.1f}% | Doc: {interim_metrics.get('doc_acc', 0):.1f}%"
            )

            if stopped_by_token_cap:
                break

            if batch_index < len(batches):
                await asyncio.sleep(BATCH_GAP_SEC)

    subtest_elapsed_sec = time.perf_counter() - subtest_start
    progress.finish()

    return results, stopped_by_token_cap, subtest_elapsed_sec


async def run_difficulty(
        difficulty: str,
        limiter: SlidingWindowRateLimiter
) -> Dict[str, Any]:
    requests_by_concurrency = build_requests_for_difficulty(difficulty)
    token_guard = TokenBudgetGuard(TOKEN_CAP_PER_DIFFICULTY_RUN)

    overall_results: List[EvalResult] = []
    stopped_early = False
    subtest_summaries: List[Dict[str, Any]] = []

    for concurrency in CONCURRENCY_LEVELS:
        sub_results, token_stop, sub_elapsed_sec = await run_subtest(
            difficulty=difficulty,
            concurrency=concurrency,
            requests=requests_by_concurrency[concurrency],
            limiter=limiter,
            token_guard=token_guard,
        )
        overall_results.extend(sub_results)

        metrics = compute_metrics(sub_results)
        subtest_summaries.append({
            "difficulty": difficulty,
            "concurrency": concurrency,
            "requests_completed": len(sub_results),
            "total_tokens": metrics.get("total_tokens", 0),
            "success_rate_pct": metrics.get("success_rate_pct", 0.0),
            "routing_accuracy_pct": metrics.get("routing_accuracy_pct", 0.0),
            "format_compliance_pct": metrics.get("format_compliance_pct", 0.0),
            "latency_lt_500_pct": metrics.get("latency_lt_500_pct", 0.0),
            "state_retention_pct": metrics.get("state_retention_pct", "N/A"),
            "p50_latency_ms": metrics.get("p50_latency_ms", 0),
            "p95_latency_ms": metrics.get("p95_latency_ms", 0),
            "subtest_elapsed_sec": sub_elapsed_sec,
            "chat_acc": metrics.get("chat_acc", 0.0),
            "task_acc": metrics.get("task_acc", 0.0),
            "dev_acc": metrics.get("dev_acc", 0.0),
            "doc_acc": metrics.get("doc_acc", 0.0),
            "stopped_by_token_cap": token_stop,
        })

        if token_stop:
            stopped_early = True
            print(
                f"Token cap reached during {difficulty} at concurrency {concurrency}. Remaining subtests will be skipped.")
            break

    difficulty_metrics = compute_metrics(overall_results)

    return {
        "difficulty": difficulty,
        "metrics": difficulty_metrics,
        "subtest_summary": subtest_summaries,
        "stopped_early": stopped_early,
        "requests_completed": len(overall_results),
        "tokens_consumed": token_guard.total_tokens,
        "results": overall_results,
    }


# =========================================================
# ENTRY POINT
# =========================================================
def choose_difficulties() -> List[str]:
    print("Select test scope:")
    print("1) Simple")
    print("2) Complex/Interrupt")
    print("3) Edge Case")
    print("4) All")
    choice = input("Enter choice [1-4]: ").strip()
    if choice == "1":
        return ["Simple"]
    if choice == "2":
        return ["Complex_Interrupt"]
    if choice == "3":
        return ["Edge_Case"]
    if choice == "4":
        return ["Simple", "Complex_Interrupt", "Edge_Case"]
    raise ValueError("Invalid choice. Please enter 1, 2, 3, or 4.")


async def async_main() -> None:
    if API_KEY == "PASTE_YOUR_GEMINI_API_KEY_HERE":
        print("\n[🚨 严重错误] 请先在代码第 25 行填入你真实的 API_KEY！")
        return

    selected = choose_difficulties()
    planned_requests = len(selected) * REQUESTS_PER_DIFFICULTY
    print(f"Planned requests for this run: {planned_requests}")
    print(f"Daily request budget assumption: {DAILY_REQUEST_BUDGET}")
    print(f"Token cap per difficulty run: {TOKEN_CAP_PER_DIFFICULTY_RUN}")

    if planned_requests > DAILY_REQUEST_BUDGET:
        print("WARNING: planned requests exceed the stated daily request budget.")
        answer = input("Continue anyway? [y/N]: ").strip().lower()
        if answer != "y":
            print("Aborted by user.")
            return

    run_root = ROOT_DIR / f"run_adversarial_minimal_{timestamp_str()}"
    ensure_dir(run_root)
    limiter = SlidingWindowRateLimiter(GLOBAL_RPM)

    difficulty_reports: List[Dict[str, Any]] = []
    all_results: List[EvalResult] = []

    for difficulty in selected:
        print(f"\nStarting difficulty: {difficulty}")
        rep = await run_difficulty(difficulty, limiter)
        difficulty_reports.append(rep)
        all_results.extend(rep["results"])

    run_config = {
        "model": MODEL_NAME,
        "max_output_tokens": MAX_OUTPUT_TOKENS,
        "temperature": TEMPERATURE,
        "top_p": TOP_P,
        "difficulties_tested": ", ".join(selected),
        "total_requests_evaluated": len(all_results),
        "requests_per_subtest": REQUESTS_PER_SUBTEST,
        "concurrency_levels": str(CONCURRENCY_LEVELS),
        "batch_gap_after_each_completed_batch_(sec)": BATCH_GAP_SEC,
        "global_RPM_limiter": GLOBAL_RPM,
        "request_timeout_(sec)": REQUEST_TIMEOUT_SEC,
        "max_retries": MAX_RETRIES,
        "daily_request_budget_assumption": DAILY_REQUEST_BUDGET,
        "token_cap_per_difficulty_run": TOKEN_CAP_PER_DIFFICULTY_RUN,
        "complexity_amplification_mode": f"cubic ({COMPLEXITY_POWER} stacked prompt layers)",
        "thinking_level": "Minimal (Compute Limited)",
        "random_seed": SEED,
    }

    report_path = run_root / "Adversarial_Benchmark_Minimal.docx"
    generate_comprehensive_word_report(
        report_path,
        run_config,
        all_results,
        difficulty_reports,
    )

    print(f"\nRun completed. Output saved to: {run_root}")
    print(f"Generated file:")
    print(f" - {report_path}")


def main() -> None:
    asyncio.run(async_main())


if __name__ == "__main__":
    main()